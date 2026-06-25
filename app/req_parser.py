"""
Requirement table parser for DOCX specifications.
Extracts individual requirement rows from Stellantis CTS tables
for precise, requirement-instance-level validation.
"""
import re
from typing import List, Dict, Any, Optional
from docx import Document


# Patterns to identify requirement table columns
REQ_ID_PATTERNS = [
    # "Requirement no. (v)", "Requirement number", "Req ID", "Req #", "Req numéro"
    re.compile(r'req(?:uirement)?\s*(?:no\.?|number|num(?:éro)?|id|#)', re.IGNORECASE),
    # "Number of the requirement", "No. of req", "Numéro de la req"
    re.compile(r'(?:no\.?|number|num(?:éro)?)\s*(?:of\s*)?(?:the\s*)?req(?:uirement)?', re.IGNORECASE),
    # Standalone "Requirement" / "Req" (lone column header)
    re.compile(r'^req(?:uirement)?\s*$', re.IGNORECASE),
    # "Identifier", "Identifiant"
    re.compile(r'identif', re.IGNORECASE),
    # Standalone "No." / "N°" / "Nº" as a column header
    re.compile(r'^(?:no\.?|n[°º])\s*(?:\([^)]*\))?\s*$', re.IGNORECASE),
]
DESC_PATTERNS = [
    re.compile(r'desc(?:ription)?', re.IGNORECASE),
    re.compile(r'(?:functional\s*)?req(?:uirement)?\s*(?:description|text)', re.IGNORECASE),
    re.compile(r'^description$', re.IGNORECASE),
]
INPUT_REQ_PATTERNS = [
    re.compile(r'input\s*req(?:uirement)?', re.IGNORECASE),
    re.compile(r'upstream\s*req(?:uirement)?', re.IGNORECASE),
    re.compile(r'(?:amont|entrante)', re.IGNORECASE),
]
# [COMMENTED OUT — Validation Plans / Test Methods — can be re-enabled later]
# VALIDATION_PATTERNS = [
#     re.compile(r'valid(?:ation)?\s*(?:method|criteria|test)?', re.IGNORECASE),
#     re.compile(r'acceptance\s*criteria', re.IGNORECASE),
#     re.compile(r'test\s*(?:method|condition)', re.IGNORECASE),
# ]
VALIDATION_PATTERNS = []  # Disabled — no validation column detection for now


def _match_column(header_text: str, patterns: List[re.Pattern]) -> bool:
    """Check if a column header matches any of the given patterns."""
    for p in patterns:
        if p.search(header_text):
            return True
    return False


def _identify_columns(headers: List[str]) -> Dict[str, int]:
    """
    Identify which column index corresponds to which field.
    Returns dict mapping field_name -> column_index.
    """
    mapping: Dict[str, int] = {}
    for i, h in enumerate(headers):
        if _match_column(h, REQ_ID_PATTERNS):
            mapping["req_id"] = i
        elif _match_column(h, DESC_PATTERNS):
            mapping["description"] = i
        elif _match_column(h, INPUT_REQ_PATTERNS):
            mapping["input_requirement"] = i
        # [COMMENTED OUT — Validation column detection disabled]
        # elif _match_column(h, VALIDATION_PATTERNS):
        #     mapping["validation"] = i
    return mapping


def extract_requirements_from_docx(file_path: str) -> List[Dict[str, Any]]:
    """
    Extract individual requirement rows from a Stellantis CTS document.
    
    Detects requirement tables by scanning for characteristic column headers
    (Requirement Number, Description, Input Requirement, Validation Method).
    Returns a list of requirement dicts with structured fields.

    Args:
        file_path: Path to the .docx file

    Returns:
        List of requirement dicts: [{req_id, description, input_requirement, 
                                     validation, section_context, table_index, row_index}]
    """
    doc = Document(file_path)
    requirements: List[Dict[str, Any]] = []
    section_context = ""
    section_stack: List[str] = []

    def update_section(heading_text: str, level: int):
        nonlocal section_stack, section_context
        section_stack = section_stack[:level - 1] if level > 1 else []
        section_stack.append(heading_text)
        section_context = " > ".join(section_stack)

    # ── CRITICAL FIX: Interleave paragraphs AND tables in document order ──
    # Previously: Phase 1 read ALL paragraphs, Phase 2 read ALL tables.
    # Bug: section_context was frozen at the LAST heading in the document,
    #       so ALL tables got the same (wrong) section.
    # Fix: Walk the XML body elements in order (w:p and w:tbl interleaved).
    
    # Build a lookup from table XML element to table index
    tbl_to_idx = {}
    for i, table in enumerate(doc.tables):
        tbl_to_idx[table._tbl] = i
    
    # Track which heading applies to each table by walking body elements
    table_section_map: Dict[int, str] = {}  # table_idx -> section_context
    
    from docx.oxml.ns import qn
    body = doc.element.body
    
    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        
        if tag == 'p':  # Paragraph
            # Check if it's a heading by looking at pStyle
            pPr = child.find(qn('w:pPr'))
            if pPr is not None:
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None:
                    style_val = pStyle.get(qn('w:val'), '')
                    if style_val.startswith('Heading') or style_val.startswith('Titre') or style_val.startswith('heading'):
                        try:
                            level = int(''.join(filter(str.isdigit, style_val)) or '1')
                        except ValueError:
                            level = 1
                        # Get paragraph text
                        text_parts = []
                        for t in child.iter(qn('w:t')):
                            if t.text:
                                text_parts.append(t.text)
                        heading_text = ''.join(text_parts).strip()
                        if heading_text:
                            update_section(heading_text, min(level, 6))
        
        elif tag == 'tbl':  # Table
            if child in tbl_to_idx:
                table_idx = tbl_to_idx[child]
                table_section_map[table_idx] = section_context
    
    # If no tables were mapped via XML walk (fallback for simple docs),
    # use the last known section_context for all tables
    if not table_section_map:
        for i in range(len(doc.tables)):
            table_section_map[i] = section_context

    # Phase 2: Scan all tables for requirement tables
    # ── Strict table type detection ──────────────────────────────────
    # Table types found in Stellantis CTS documents:
    #   TYPE A (REQUIREMENT):  Requirement Number | Description | Input requirement | Validation
    #   TYPE B (REFERENCE):    Mark | Reference | Version | Title  ← SKIP these
    #   TYPE C (DOC INFO):     Author | Date | Signature | Index | Modification  ← SKIP
    #   TYPE D (DATA/CONFIG):  Parameter | Value | Unit  ← SKIP (1-2 cols, no req ID)
    #   TYPE E (VARIANT):      Complexity criteria | Variants  ← SKIP
    
    # Negative patterns — if ANY of these match, it's NOT a requirement table
    REFERENCE_TABLE_COLUMNS = {
        "mark", "reference", "version", "title", "link", "document",
        "author", "date", "signature", "inspector", "approved", "co-author",
        "index", "modification", "status", "complexity", "variant",
    }
    
    for table_idx, table in enumerate(doc.tables):
        if len(table.rows) < 2:
            continue

        # Use the section context that was active when this table appeared
        current_sec = table_section_map.get(table_idx, section_context)

        # Read header row
        header_cells = [cell.text.strip() for cell in table.rows[0].cells]
        if not header_cells:
            continue
        
        header_lower = " | ".join(header_cells).lower()
        
        # ── RULE 1: Skip reference tables (Mark | Reference | Version | Title) ──
        # If the table has a "Mark" column, it's a reference/doc table, NOT requirements
        has_mark_column = any(
            re.match(r'^\s*mark\s*$', h, re.IGNORECASE) for h in header_cells
        )
        if has_mark_column:
            continue  # DEFINITELY a reference table — skip
        
        # ── RULE 2: Skip document info tables ──
        metadata_hits = sum(1 for m in REFERENCE_TABLE_COLUMNS if m in header_lower)
        # If 2+ metadata columns and table has ≤4 columns, it's not a requirement table
        if metadata_hits >= 2 and len(header_cells) <= 4:
            continue
        
        # ── RULE 3: Skip single-column or 2-column tables (data/config tables) ──
        if len(header_cells) <= 2:
            continue
        
        # ── RULE 4: POSITIVE check — must have a requirement-like column ──
        col_map = _identify_columns(header_cells)
        
        # A requirement table MUST have BOTH a req_id column AND a description column
        has_req_id_col = "req_id" in col_map
        has_desc_col = "description" in col_map
        has_input_col = "input_requirement" in col_map
        has_validation_col = "validation" in col_map
        
        # STRICT: Need at minimum "description" column + at least one of (req_id, input_req, validation)
        if not has_desc_col:
            continue  # No description column = not a requirement table
        
        if not (has_req_id_col or has_input_col or has_validation_col):
            # Has a description column but nothing else requirement-specific
            # This could be a generic data table — skip unless it's clearly requirements
            continue
        
        # ── RULE 5: Verify the table actually has requirement-like content ──
        # Check at least one data row for requirement patterns (SHALL/MUST, IF/THEN, etc.)
        if not has_req_id_col:
            # Without a req_id column, verify the description cells look like requirements
            sample_rows_checked = 0
            requirement_like_rows = 0
            for row in table.rows[1:min(4, len(table.rows))]:
                cells = [cell.text.strip() for cell in row.cells]
                if has_desc_col and col_map["description"] < len(cells):
                    desc_text = cells[col_map["description"]]
                    if desc_text and len(desc_text) > 10:
                        sample_rows_checked += 1
                        # Check for requirement keywords
                        if re.search(r'\b(shall|must|will|if|when|then|system|component|function)\b', 
                                    desc_text, re.IGNORECASE):
                            requirement_like_rows += 1
            
            # If we sampled rows and none look like requirements, skip this table
            if sample_rows_checked > 0 and requirement_like_rows == 0:
                continue

        # Extract requirement rows
        for row_idx, row in enumerate(table.rows[1:], 1):  # Skip header
            cells = [cell.text.strip() for cell in row.cells]
            if not any(cells):  # Skip empty rows
                continue
            
            req: Dict[str, Any] = {
                "table_index": table_idx,
                "row_index": row_idx,
                "section_context": current_sec,
            }
            
            if "req_id" in col_map and col_map["req_id"] < len(cells):
                raw_id = cells[col_map["req_id"]]
                # ── Clean up multi-line ID cells ──────────────────────
                # Some templates put ASIL attributes, PSA comments, etc.
                # in the same cell as the requirement ID, separated by newlines.
                # Example cell content:
                #   REF-PSP-AIRBAG-FRONT-001
                #   Att_Sdf@ ASIL_A(A)
                #   PSA_Comments@{{if you want to add some information}}
                # We extract ONLY the first line as the actual requirement ID.
                # Also strip anything after '@' (annotation marker) on the first line.
                if '\n' in raw_id:
                    raw_id = raw_id.split('\n')[0].strip()
                if '@' in raw_id:
                    # Annotation format like "Att_Sdf@ ASIL_A(A)" — keep only the ID part
                    raw_id = raw_id.split('@')[0].strip()
                req["req_id"] = raw_id
            if "description" in col_map and col_map["description"] < len(cells):
                req["description"] = cells[col_map["description"]]
            if "input_requirement" in col_map and col_map["input_requirement"] < len(cells):
                req["input_requirement"] = cells[col_map["input_requirement"]]
            # [COMMENTED OUT — Validation column extraction disabled]
            # if "validation" in col_map and col_map["validation"] < len(cells):
            #     req["validation"] = cells[col_map["validation"]]

            # Only keep rows that have at least a description
            if req.get("description") or req.get("req_id"):
                requirements.append(req)

    return requirements


def format_requirements_for_prompt(requirements: List[Dict[str, Any]], max_items: int = 50) -> str:
    """
    Format extracted requirements as a structured summary for the LLM prompt.
    """
    if not requirements:
        return "[No structured requirement rows detected in the document]"

    lines = [f"═══ EXTRACTED REQUIREMENT ROWS ({len(requirements)} total, showing first {min(len(requirements), max_items)}) ═══"]
    
    # Group by section
    by_section: Dict[str, List[Dict]] = {}
    for r in requirements:
        sec = r.get("section_context", "Unknown section")
        by_section.setdefault(sec, []).append(r)

    for section, reqs in by_section.items():
        lines.append(f"\n--- {section} ({len(reqs)} requirements) ---")
        for r in reqs[:max_items]:
            req_id = r.get("req_id", "?")
            desc = (r.get("description", "") or "")[:150]
            input_req = (r.get("input_requirement", "") or "")[:80]
            validation = (r.get("validation", "") or "")[:80]
            
            parts = [f"  [{req_id}] {desc}"]
            if input_req and input_req not in ("N/A", "n/a", ""):
                parts.append(f"    Input: {input_req}")
            if validation:
                parts.append(f"    Validation: {validation}")
            lines.append("\n".join(parts))

    return "\n".join(lines)


def analyze_requirement_quality(requirements: List[Dict[str, Any]]) -> Dict[str, Any]:
    """
    Quick statistical analysis of requirement quality from structured data.
    Returns counts that can inform the LLM prompt.
    """
    total = len(requirements)
    if total == 0:
        return {"total": 0}

    with_id = sum(1 for r in requirements if r.get("req_id"))
    with_input = sum(1 for r in requirements if r.get("input_requirement") and 
                     r["input_requirement"] not in ("N/A", "n/a", "", "N / A"))
    with_validation = sum(1 for r in requirements if r.get("validation"))
    
    # Detect potentially weak descriptions (short, vague)
    weak_descs = 0
    vague_patterns = [r'\bTBD\b', r'\bXXX\b', r'<<.*?>>', r'\bnote\b', r'\bto be\b']
    for r in requirements:
        desc = r.get("description", "")
        if not desc:
            weak_descs += 1
        elif len(desc) < 30:
            weak_descs += 1
        else:
            for pat in vague_patterns:
                if re.search(pat, desc, re.IGNORECASE):
                    weak_descs += 1
                    break

    return {
        "total": total,
        "with_requirement_id": with_id,
        "with_input_requirement": with_input,
        "with_validation_method": with_validation,
        "potentially_weak_descriptions": weak_descs,
        "requirement_id_coverage": round(with_id / total * 100, 1) if total else 0,
        "input_traceability_coverage": round(with_input / total * 100, 1) if total else 0,
    }
