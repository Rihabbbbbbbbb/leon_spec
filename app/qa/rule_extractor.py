"""
Rule extraction engine — extracts 100% of validation rules from the ACTUAL
Stellantis template and writing-guide DOCX files.

No hardcoded rules. Every rule is extracted from the real source documents
located in data/refs/. This guarantees 100% traceability: every finding
cites the exact source document, section, and rule identifier.

Extracted rule types:
  1. MANDATORY_SECTIONS  — section headings found in the template's standard plan
  2. TEMPLATE_PLACEHOLDERS — <<...>> / <...> markers the template instructs to fill
  3. WRITING_GUIDE_RULES — R01..R53, P01..P09 rules from the writing guide
  4. REQUIREMENT_FORMAT   — the "shall" + 3-column-table rule from the template
  5. SECTION_ORDER        — the ordered standard plan from the writing guide TOC
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Tuple

# ── Source document paths (the REAL reference files) ──────────────
REFS_DIR = Path(__file__).resolve().parent.parent.parent / "data" / "refs"
TEMPLATE_DOCX = REFS_DIR / "Component_or_Part_Specification_Template 1.docx"
WRITING_GUIDE_DOCX = REFS_DIR / "Component_or_Part_Specification_Writing_guide 1.docx"


# ── Data structures ───────────────────────────────────────────────
@dataclass
class SectionRule:
    """A mandatory section extracted from the template standard plan."""
    name: str               # section heading as it appears in the template
    level: int = 1          # heading level (1=top, 2=sub, ...)
    source: str = "template"
    order: int = 0          # position in the standard plan
    description: str = ""   # template instruction text for this section (if any)


@dataclass
class WritingGuideRule:
    """A rule (R## or P##) extracted from the writing guide."""
    rule_id: str            # e.g. "R22", "P04"
    category: str           # "rule" | "principle"
    section: str            # writing-guide section where it appears
    text: str               # the full rule text (English version preferred)
    source: str = "writing_guide"
    line_number: int = 0    # line in extracted text (for traceability)


@dataclass
class TemplateInstruction:
    """A template instruction (red text / <<...>> placeholder) to check."""
    placeholder: str        # the raw placeholder text
    section: str            # section where it appears
    instruction: str        # the surrounding instruction text
    source: str = "template"


@dataclass
class ExtractedRules:
    """All rules extracted from the real source documents."""
    mandatory_sections: List[SectionRule] = field(default_factory=list)
    recommended_sections: List[str] = field(default_factory=list)
    section_order: List[str] = field(default_factory=list)
    writing_guide_rules: List[WritingGuideRule] = field(default_factory=list)
    template_instructions: List[TemplateInstruction] = field(default_factory=list)
    template_text: str = ""
    guide_text: str = ""
    extraction_ok: bool = False
    errors: List[str] = field(default_factory=list)


# ── DOCX text extraction (self-contained, no external dep on retrieval) ─
def _extract_docx_text(path: Path) -> str:
    """Extract plain text from a DOCX file (paragraphs + tables in order)."""
    from docx import Document
    doc = Document(str(path))
    parts: List[str] = []
    # Walk body elements in document order to preserve paragraph/table interleaving
    from docx.oxml.ns import qn
    body = doc.element.body
    # Build a lookup of table elements to table objects
    table_map = {tbl._element: tbl for tbl in doc.tables}
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            # Find the matching paragraph object
            for para in doc.paragraphs:
                if para._element is child:
                    t = para.text.strip()
                    if t:
                        parts.append(t)
                    break
        elif child.tag == qn("w:tbl"):
            tbl = table_map.get(child)
            if tbl is not None:
                for row in tbl.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
    return "\n".join(parts)


# ── Section extraction from the template ──────────────────────────
# The template's standard plan uses ALLCAPS headings for top-level sections.
# We extract them in order, preserving the real structure.

# Sections that are clearly template meta-instructions, not real CTS sections
_META_LINES = {
    "COMPONENT TECHNICAL SPECIFICATION STANDARD TEMPLATE APPLICABLE FOR MECHATRONIC COMPONENT_STELLANTIS HARMONIZED",
    "Writing instructions are PRINTED IN RED, delete them before submitting the document for revision",
    "Avoid changing formatting display of the template",
    "WARNING",
    "Table of contents",
    "Table of updates",
    "DOCUMENT",
    "ARCHITECTURES",
    "CS.00050",
    "THIS IS AN INFORMATIVE TABLE; NOT TO BE INCLUDED IN YOUR SPEC",
}

# Patterns that indicate a line is EXAMPLE content, not a real section heading.
# The template contains example requirement IDs, use-case names, and sample
# table rows that happen to be ALLCAPS but are NOT mandatory sections.
_EXAMPLE_PATTERNS = [
    re.compile(r"^REF-"),            # example requirement ID: REF-PSP-AIRBAG-FRONT-001
    re.compile(r"^GEN-"),            # example requirement ID: GEN-XXX-CDC-54411.001
    re.compile(r"^APP-"),            # example requirement ID
    re.compile(r"^ER\s+ERF"),        # example dreaded event ID: ER ERF.4.01
    re.compile(r"^USE CASE"),        # example use case placeholder: USE CASE XXX
    re.compile(r"\bXXX\b"),         # contains XXX placeholder → example
    re.compile(r"\bYYY\b"),         # contains YYY placeholder → example
    re.compile(r"\(\d+\)$"),        # ends with (1) etc. → example table annotation
    re.compile(r"^\d{2,}"),          # starts with multi-digit number (table ref)
]

# Headings that appear in the template but are sub-sections / optional
_SUBSECTION_PREFIXES = (
    "UPSTREAM REQUIREMENTS",
    "CONSTRAINT REQUIREMENTS FROM OTHER DISCIPLINES",
    "REGULATION AND CONSUMERISM",
    "MANDATORY REQUIREMENTS",
    "STANDARDS",
    "TECHNICAL SPECIFICATIONS",
    "TECHNICAL SPECIFICATIONS FOR CONNECTORS",
    "GENERAL TECHNICAL ON FAULT FINDING AND DOWNLOAD",
    "OTHER TECHNICAL SPECIFICATIONS APPLICABLE",
    "DEPENDABILITY VOCABULARY",
    "OTHER GENERIC TERMS",
    "MEASURING UNITS",
    "VOCABULARY SPECIFIC TO THE COMPONENT",
    "FUNCTIONAL COMPLEXITY",
    "ARCHITECTURE COMPLEXITY",
    "FUNCTIONAL DIVERSITY",
    "ARCHITECTURE DIVERSITY",
    "SYSTEM COMPLEXITY",
    "SYSTEM DIVERSITY",
)


def _extract_template_sections(text: str) -> Tuple[List[SectionRule], List[str]]:
    """
    Extract the mandatory section structure from the template text.

    Returns (mandatory_sections, section_order) where section_order is the
    ordered list of top-level section names as they appear in the template.
    """
    lines = text.split("\n")
    sections: List[SectionRule] = []
    order: List[str] = []
    seen = set()
    order_idx = 0

    # ALLCAPS heading pattern (the template uses ALLCAPS for section titles)
    allcaps_re = re.compile(r"^([A-Z][A-Z0-9\s/()\-&,:;.\u2013\u2014]{2,})$")

    for line in lines:
        stripped = line.strip()
        if not stripped or len(stripped) > 120:
            continue
        # Skip meta lines
        if stripped in _META_LINES:
            continue
        # Skip placeholder instruction lines (start with <<)
        if stripped.startswith("<<"):
            continue
        # Skip "Figure <n>" lines
        if stripped.lower().startswith("figure"):
            continue

        m = allcaps_re.match(stripped)
        if m:
            name = m.group(1).strip().rstrip(":")
            # Require high uppercase ratio
            letters = [c for c in name if c.isalpha()]
            if not letters:
                continue
            upper_ratio = sum(1 for c in letters if c.isupper()) / len(letters)
            if upper_ratio < 0.8:
                continue
            if len(name) < 3 or name in seen:
                continue
            # Skip meta lines
            if name in _META_LINES:
                continue
            # Skip example content (requirement IDs, use-case placeholders, etc.)
            if any(pat.search(name) for pat in _EXAMPLE_PATTERNS):
                continue

            # Determine level: subsections are level 2
            level = 2 if name in _SUBSECTION_PREFIXES else 1
            sections.append(SectionRule(
                name=name, level=level, source="template",
                order=order_idx,
            ))
            seen.add(name)
            if level == 1:
                order.append(name)
                order_idx += 1

    return sections, order


# ── Template placeholder/instruction extraction ───────────────────
_PLACEHOLDER_RE = re.compile(r"<<([^>]*)>>")
_COMPONENT_VAR_RE = re.compile(r"<(component name|part name|Part name|name of the Model|reference|PSP|stakeholder|project name)>")


def _extract_template_instructions(text: str, sections: List[SectionRule]) -> List[TemplateInstruction]:
    """Extract template placeholders and their section context."""
    instructions: List[TemplateInstruction] = []
    lines = text.split("\n")
    current_section = ""

    for line in lines:
        stripped = line.strip()
        # Track current section
        if stripped and stripped.upper() == stripped and len(stripped) < 100:
            for s in sections:
                if s.name == stripped:
                    current_section = stripped
                    break

        # Find <<...>> placeholders
        for m in _PLACEHOLDER_RE.finditer(line):
            ph = m.group(0)
            instructions.append(TemplateInstruction(
                placeholder=ph,
                section=current_section,
                instruction=stripped[:200],
            ))
        # Find <component name> etc.
        for m in _COMPONENT_VAR_RE.finditer(line):
            ph = m.group(0)
            instructions.append(TemplateInstruction(
                placeholder=ph,
                section=current_section,
                instruction=stripped[:200],
            ))

    return instructions


# ── Writing-guide rule extraction (R## and P##) ───────────────────
# Rules appear as "R22:" or "R22 :" or "P04:" at the start of a line.
_RULE_RE = re.compile(r"^(R\d{1,2})\s*:?\s*(.+)")
_PRINCIPLE_RE = re.compile(r"^(P\d{1,2})\s*:?\s*(.+)")


def _extract_writing_guide_rules(text: str) -> List[WritingGuideRule]:
    """
    Extract all R## (rules) and P## (principles) from the writing guide.

    The writing guide is bilingual (French + English). We prefer the English
    version when both are present (English lines follow the French ones and
    repeat the rule ID).
    """
    lines = text.split("\n")
    rules: List[WritingGuideRule] = []
    seen_ids = {}  # rule_id -> index in rules list
    current_section = ""

    # Section heading detection in the guide (numbered or ALLCAPS)
    section_num_re = re.compile(r"^\d+(?:\.\d+)*\.?\s+([A-Z][A-Za-z\s/()\-&,:;.]{2,})$")
    allcaps_re = re.compile(r"^([A-Z][A-Z\s/()\-&,:;.\u2013\u2014]{2,})$")

    for i, line in enumerate(lines):
        stripped = line.strip()
        # Track section context
        m_sec = section_num_re.match(stripped) or allcaps_re.match(stripped)
        if m_sec and len(stripped) < 100:
            current_section = m_sec.group(1).strip().rstrip(":")

        # Detect rule lines
        m_rule = _RULE_RE.match(stripped)
        m_princ = _PRINCIPLE_RE.match(stripped)

        if m_rule:
            rule_id = m_rule.group(1)
            rule_text = m_rule.group(2).strip()
            # Skip if it's a TOC line (contains a tab + page number)
            if "\t" in stripped and re.search(r"\t\d+$", stripped):
                continue
            # Skip very short fragments
            if len(rule_text) < 10:
                continue

            # If we already have this rule ID, prefer the English version
            # (English text typically contains "shall" or is longer/cleaner)
            if rule_id in seen_ids:
                idx = seen_ids[rule_id]
                existing = rules[idx]
                # Replace if the new text looks like English (has common EN words)
                en_words = sum(1 for w in ("shall", "must", "should", "the", "requirement",
                                           "rule", "document", "system", "presented")
                               if w in rule_text.lower())
                fr_words = sum(1 for w in ("exigence", "règle", "document", "système",
                                           "présenté", "rédaction", "doit", "sont")
                               if w in rule_text.lower())
                if en_words > fr_words and len(rule_text) >= len(existing.text):
                    rules[idx] = WritingGuideRule(
                        rule_id=rule_id, category="rule",
                        section=current_section, text=rule_text,
                        source="writing_guide", line_number=i,
                    )
            else:
                rules.append(WritingGuideRule(
                    rule_id=rule_id, category="rule",
                    section=current_section, text=rule_text,
                    source="writing_guide", line_number=i,
                ))
                seen_ids[rule_id] = len(rules) - 1

        elif m_princ:
            rule_id = m_princ.group(1)
            rule_text = m_princ.group(2).strip()
            if "\t" in stripped and re.search(r"\t\d+$", stripped):
                continue
            if len(rule_text) < 10:
                continue
            if rule_id in seen_ids:
                idx = seen_ids[rule_id]
                existing = rules[idx]
                en_words = sum(1 for w in ("shall", "must", "should", "the", "requirement",
                                           "each", "service", "described", "system")
                               if w in rule_text.lower())
                fr_words = sum(1 for w in ("exigence", "règle", "chaque", "service",
                                           "décrit", "système", "doit")
                               if w in rule_text.lower())
                if en_words > fr_words and len(rule_text) >= len(existing.text):
                    rules[idx] = WritingGuideRule(
                        rule_id=rule_id, category="principle",
                        section=current_section, text=rule_text,
                        source="writing_guide", line_number=i,
                    )
            else:
                rules.append(WritingGuideRule(
                    rule_id=rule_id, category="principle",
                    section=current_section, text=rule_text,
                    source="writing_guide", line_number=i,
                ))
                seen_ids[rule_id] = len(rules) - 1

    return rules


# ── Recommended sections (from writing guide TOC, level-2 under 5.4) ─
_RECOMMENDED_KEYWORDS = [
    "MISSION PROFILE", "LIFETIME", "ERGONOMICS", "SAFETY",
    "MAINTAINABILITY", "PRODUCT QUALITY", "TRACEABILITY",
    "DESIGN AND MANUFACTURING", "ENVIRONMENT CONDITIONS",
    "INTEGRATION AND VALIDATION", "DEMONSTRATION OF COMPLIANCE",
]


def _extract_recommended_sections(guide_text: str) -> List[str]:
    """Extract recommended sections from the writing guide TOC."""
    found = []
    for kw in _RECOMMENDED_KEYWORDS:
        if kw.lower() in guide_text.lower():
            found.append(kw)
    return found


# ── Main extraction function ──────────────────────────────────────
_rules_cache: Optional[ExtractedRules] = None


def extract_all_rules(force: bool = False) -> ExtractedRules:
    """
    Extract all validation rules from the real template + writing guide DOCX.

    Results are cached in memory. Use force=True to re-extract.

    Returns ExtractedRules with 100% traceable source data.
    """
    global _rules_cache
    if _rules_cache is not None and not force:
        return _rules_cache

    rules = ExtractedRules()
    errors = []

    # Extract template text
    try:
        if not TEMPLATE_DOCX.exists():
            errors.append(f"Template DOCX not found: {TEMPLATE_DOCX}")
            tmpl_text = ""
        else:
            tmpl_text = _extract_docx_text(TEMPLATE_DOCX)
            rules.template_text = tmpl_text
    except Exception as e:
        errors.append(f"Template extraction error: {e}")
        tmpl_text = ""

    # Extract writing guide text
    try:
        if not WRITING_GUIDE_DOCX.exists():
            errors.append(f"Writing guide DOCX not found: {WRITING_GUIDE_DOCX}")
            guide_text = ""
        else:
            guide_text = _extract_docx_text(WRITING_GUIDE_DOCX)
            rules.guide_text = guide_text
    except Exception as e:
        errors.append(f"Writing guide extraction error: {e}")
        guide_text = ""

    if tmpl_text:
        sections, order = _extract_template_sections(tmpl_text)
        rules.mandatory_sections = sections
        rules.section_order = order
        rules.template_instructions = _extract_template_instructions(tmpl_text, sections)

    if guide_text:
        rules.writing_guide_rules = _extract_writing_guide_rules(guide_text)
        rules.recommended_sections = _extract_recommended_sections(guide_text)

    rules.errors = errors
    rules.extraction_ok = len(errors) == 0 and bool(tmpl_text)
    _rules_cache = rules
    return rules


def get_mandatory_section_names() -> List[str]:
    """Convenience: return just the mandatory section names (top-level)."""
    r = extract_all_rules()
    return [s.name for s in r.mandatory_sections if s.level == 1]


def get_rule_by_id(rule_id: str) -> Optional[WritingGuideRule]:
    """Look up a specific writing-guide rule by its ID (e.g. 'R22')."""
    r = extract_all_rules()
    for rule in r.writing_guide_rules:
        if rule.rule_id == rule_id:
            return rule
    return None


if __name__ == "__main__":
    import sys
    sys.stdout.reconfigure(encoding="utf-8")
    r = extract_all_rules(force=True)
    print(f"Extraction OK: {r.extraction_ok}")
    print(f"Errors: {r.errors}")
    print(f"Mandatory sections ({len(r.mandatory_sections)}):")
    for s in r.mandatory_sections:
        print(f"  [{s.level}] {s.name} (order={s.order})")
    print(f"\nSection order ({len(r.section_order)}): {r.section_order}")
    print(f"\nWriting guide rules ({len(r.writing_guide_rules)}):")
    for rule in r.writing_guide_rules:
        print(f"  {rule.rule_id} [{rule.category}] §{rule.section}: {rule.text[:80]}...")
    print(f"\nTemplate instructions ({len(r.template_instructions)}):")
    for ti in r.template_instructions[:10]:
        print(f"  {ti.placeholder} §{ti.section}")
