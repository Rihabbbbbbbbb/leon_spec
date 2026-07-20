"""
Unified DOCX report generator for specification validation.

This is the STANDARDIZED TEMPLATE that LEON always uses to deliver
specification validation reports to the user as a downloadable document.

Designed to be COMPACT and ENGINEER-FOCUSED (target: 3-6 pages):
  1. Synthèse       — verdict banner, key metadata, scores per axis, counts
  2. Problèmes à corriger — errors then warnings, each with location,
                      evidence excerpt and suggested fix (the actionable core)
  3. Couverture des sections — missing sections highlighted, found inline
  4. Recommandations — top prioritized actions
  5. Périmètre de l'analyse — rules checked, sources, method (audit trail)

Uses python-docx (already in requirements.txt) for Azure Function compatibility.
"""
from __future__ import annotations

import io
import datetime
from typing import Dict, List


# ═══════════════════════════════════════════════════════════════════
# DOCX HELPER FUNCTIONS
# ═══════════════════════════════════════════════════════════════════

def _set_cell_shading(cell, hex_color: str):
    """Apply background shading to a DOCX table cell."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _add_page_number_footer(section):
    """Add page numbers to the footer of a DOCX section."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = 1  # CENTER
    run = p.add_run("LEON — Rapport de validation | Page ")
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


_VERDICT_LABELS = {
    "GOOD": "GOOD — Spécification conforme",
    "ACCEPTABLE_WITH_FIXES": "ACCEPTABLE — corrections à prévoir",
    "NOT_RELIABLE": "NOT RELIABLE — révision approfondie nécessaire",
    "NON_COMPLIANT": "NON COMPLIANT — non conforme au template CTS",
}


def _verdict_color(verdict: str) -> tuple:
    """Return RGB color for a verdict."""
    return {
        "GOOD": (40, 167, 69),
        "ACCEPTABLE_WITH_FIXES": (176, 122, 0),
        "NOT_RELIABLE": (204, 85, 0),
        "NON_COMPLIANT": (220, 53, 69),
    }.get(verdict, (108, 117, 125))


def _verdict_hex(verdict: str) -> str:
    """Return hex color for a verdict (for cell shading)."""
    return {
        "GOOD": "C6EFCE",
        "ACCEPTABLE_WITH_FIXES": "FFEB9C",
        "NOT_RELIABLE": "FFD580",
        "NON_COMPLIANT": "FFC7CE",
    }.get(verdict, "D9D9D9")


def _severity_hex(severity: str) -> str:
    """Return hex color for a severity level (for cell shading)."""
    return {
        "error": "FFC7CE",
        "warning": "FFEB9C",
        "pass": "C6EFCE",
        "info": "D9E2F3",
    }.get(severity, "FFFFFF")


def _severity_rgb(severity: str) -> tuple:
    """Return RGB color for a severity level."""
    return {
        "error": (220, 53, 69),
        "warning": (176, 122, 0),
        "pass": (40, 167, 69),
        "info": (23, 162, 184),
    }.get(severity, (108, 117, 125))


def _score_hex(score: float) -> str:
    """Return hex color for a score value (0-1)."""
    if score >= 0.80:
        return "C6EFCE"
    elif score >= 0.60:
        return "FFEB9C"
    elif score >= 0.35:
        return "FFD580"
    else:
        return "FFC7CE"


# ═══════════════════════════════════════════════════════════════════
# MAIN DOCUMENT GENERATOR
# ═══════════════════════════════════════════════════════════════════

# Caps to keep the report short — the full machine-readable detail
# stays available in the JSON validationReport.
_MAX_PROBLEM_BLOCKS = 40      # detailed error/warning blocks
_MAX_RECOMMENDATIONS = 6
_EXCERPT_MAX = 180
_MSG_MAX = 300

# Correction examples per rule — shown in the per-problem detail blocks.
# Keyed by rule_id first, then by check category as fallback.
_RULE_EXAMPLES = {
    "TEMPLATE": 'Avant : « Response time shall be <<x>> ms » → Après : « Response time shall be 50 ms ».',
    "R20": "Format d'identifiant attendu : REF-PSP-ASU-0001 (préfixe REF-/APP-/GEN- + composant + numéro unique).",
    "R22": 'Tableau d\'exigence à 3 colonnes : « REF-ASU-CD-X-001 | The unit shall … | REQ-0508543 » (ou « N/A » si pas d\'exigence amont).',
    "R23": "Avant : « The system shall respond quickly » → Après : « The system shall respond within 100 ms ».",
    "R05": "Titre attendu : « Requirements Document of the Alarm Siren Unit (ASU) Module ».",
    "R07": "Ajouter en en-tête un tableau « Written by / Checked by / Approved by » avec noms et dates.",
    "R09": "Ajouter une « Table of updates » : Version | Date | Auteur | Nature de la modification.",
    "WRITING_GUIDE": "Ajouter la section recommandée, même brève — indiquer « Not applicable » si elle est sans objet.",
    "A_SECTION_COVERAGE": "Ajouter le titre de section manquant à sa place dans le plan standard du template, puis rédiger son contenu.",
    "E_REQUIREMENT_LANGUAGE": "Avant : « The system should log errors » → Après : « The system shall log errors ».",
    "F_REQUIREMENT_IDS": "Attribuer un identifiant unique à chaque exigence : REF-PSP-<COMPOSANT>-001, -002, …",
    "G_TRACEABILITY": "Renseigner l'exigence amont pour chaque exigence (colonne « Input requirement ») ou « N/A ».",
    "R17": "Vérifier la section APPLICABLE DOCUMENTS / STANDARDS : chaque norme citée dans une exigence (ex. [STA20], ISO 26262) doit y figurer, et chaque norme déclarée doit être réellement utilisée par au moins une exigence.",
    "J_STANDARDS_CONSISTENCY": "Vérifier la section APPLICABLE DOCUMENTS / STANDARDS : chaque norme citée dans une exigence (ex. [STA20], ISO 26262) doit y figurer, et chaque norme déclarée doit être réellement utilisée par au moins une exigence.",
}


def _example_for(finding: Dict) -> str:
    """Return a correction example for a finding, if one is defined."""
    return (
        _RULE_EXAMPLES.get(finding.get("rule_id", ""))
        or _RULE_EXAMPLES.get(finding.get("check", ""))
        or ""
    )


def generate_spec_validation_document(report: Dict) -> bytes:
    """
    Generate the standardized LEON validation report (compact DOCX).

    Args:
        report: The validation report dict from validate_with_evidence()

    Returns:
        DOCX file as bytes
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    # ── Page setup ─────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # ── Helpers ────────────────────────────────────────────
    def _add_heading(text, level=1, color=(0, 51, 102)):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
        return h

    def _add_para(text, bold=False, italic=False, size=10, color=None, align=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = RGBColor(*color)
        if align is not None:
            p.alignment = align
        return p

    def _style_header_cell(cell, size=9):
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(size)
        _set_cell_shading(cell, "003366")

    # ── Report data ────────────────────────────────────────
    file_name = report.get("fileName", "N/A")
    verdict = report.get("verdict", "UNKNOWN")
    overall_score = report.get("overallScore", 0)
    counts = report.get("summaryCounts", {})
    scores = report.get("scores", {})
    findings = report.get("findings", [])
    detailed = report.get("detailed", {})
    errors = detailed.get("errors") or [f for f in findings if f.get("severity") == "error"]
    warnings = detailed.get("warnings") or [f for f in findings if f.get("severity") == "warning"]
    sections_found = report.get("sectionsFound", [])
    sections_missing = report.get("sectionsMissing", [])
    rules_used = report.get("rulesUsed", {})

    # ═══════════════════════════════════════════════════════
    # HEADER
    # ═══════════════════════════════════════════════════════
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LEON — Rapport de Validation de Spécification")
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run(
        f"{file_name}  ·  {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}  ·  "
        f"Stellantis Mechatronics Engineering"
    )
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = RGBColor(108, 117, 125)
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════
    # 1. SYNTHÈSE — verdict banner + scores + counts
    # ═══════════════════════════════════════════════════════
    _add_heading("1. Synthèse", level=2)

    # Verdict banner (single shaded cell)
    banner = doc.add_table(rows=1, cols=1)
    banner.style = "Table Grid"
    cell = banner.cell(0, 0)
    score_txt = f"{overall_score:.0%}" if isinstance(overall_score, (int, float)) else str(overall_score)
    cell.text = f"{_VERDICT_LABELS.get(verdict, verdict)}   —   Score global : {score_txt}"
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(*_verdict_color(verdict))
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_cell_shading(cell, _verdict_hex(verdict))
    doc.add_paragraph()

    # Counts (one horizontal row)
    count_data = [
        ("Erreurs", counts.get("errors", 0), "FFC7CE"),
        ("Avertissements", counts.get("warnings", 0), "FFEB9C"),
        ("Vérifications conformes", counts.get("pass", 0), "C6EFCE"),
        ("Sections manquantes", len(sections_missing), "FFC7CE" if sections_missing else "C6EFCE"),
    ]
    counts_table = doc.add_table(rows=2, cols=len(count_data))
    counts_table.style = "Table Grid"
    counts_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, (label, value, hex_c) in enumerate(count_data):
        head = counts_table.cell(0, ci)
        head.text = label
        _style_header_cell(head)
        val_cell = counts_table.cell(1, ci)
        val_cell.text = str(value)
        val_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in val_cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(12)
        _set_cell_shading(val_cell, hex_c)
    doc.add_paragraph()

    # Scores per axis (compact)
    score_rows = [
        ("structure", "Structure (sections du template)", 0.25),
        ("section_order", "Ordre des sections", 0.05),
        ("template_cleanliness", "Propreté du template", 0.10),
        ("requirements_quality", "Qualité des exigences", 0.35),
        ("writing_guide_compliance", "Conformité au guide d'écriture", 0.25),
    ]
    score_table = doc.add_table(rows=len(score_rows) + 1, cols=3)
    score_table.style = "Table Grid"
    score_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, header in enumerate(["Axe d'évaluation", "Poids", "Score"]):
        _style_header_cell(score_table.cell(0, ci))
        score_table.cell(0, ci).text = header
        _style_header_cell(score_table.cell(0, ci))
    for ri, (key, label, weight) in enumerate(score_rows, 1):
        val = scores.get(key, 0)
        val_f = val if isinstance(val, (int, float)) else 0
        score_table.cell(ri, 0).text = label
        score_table.cell(ri, 1).text = f"{weight:.0%}"
        score_table.cell(ri, 2).text = f"{val_f:.0%}"
        hex_color = _score_hex(val_f)
        for ci in range(3):
            for run in score_table.cell(ri, ci).paragraphs[0].runs:
                run.font.size = Pt(9)
            _set_cell_shading(score_table.cell(ri, ci), hex_color)
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════
    # 2. PROBLÈMES À CORRIGER — the actionable core
    # ═══════════════════════════════════════════════════════
    n_problems = len(errors) + len(warnings)
    _add_heading(f"2. Problèmes à Corriger ({n_problems})", level=2)

    if n_problems == 0:
        _add_para(
            "✅ Aucun problème détecté — le document est conforme au template "
            "CTS et au guide d'écriture.",
            bold=True, size=11, color=(40, 167, 69)
        )
    else:
        _add_para(
            "Erreurs critiques d'abord, puis avertissements. Pour chaque "
            "problème : où il se trouve, ce qui ne va pas, et comment le corriger.",
            italic=True, size=9, color=(108, 117, 125)
        )

        problem_list = (
            [("ERREUR", f) for f in errors]
            + [("AVERTISSEMENT", f) for f in warnings]
        )
        shown = problem_list[:_MAX_PROBLEM_BLOCKS]

        prob_table = doc.add_table(rows=len(shown) + 1, cols=5)
        prob_table.style = "Table Grid"

        p_headers = ["#", "Sévérité", "Où (section / localisation)",
                     "Problème constaté", "Comment corriger"]
        for ci, header in enumerate(p_headers):
            cell = prob_table.cell(0, ci)
            cell.text = header
            _style_header_cell(cell)

        # Column widths (A4 usable ≈ 17 cm)
        col_widths = [Cm(0.9), Cm(2.2), Cm(3.4), Cm(6.0), Cm(4.5)]

        for ri, (sev_label, f) in enumerate(shown, 1):
            sev = f.get("severity", "warning")
            hex_c = _severity_hex(sev)
            where = f.get("section", "") or f.get("user_location", "") or "Document entier"
            message = (f.get("message", "") or "")[:_MSG_MAX]
            excerpt = (f.get("user_excerpt", "") or "")[:_EXCERPT_MAX]
            problem_txt = message + (f'\n« {excerpt} »' if excerpt else "")
            fix = f.get("fix_suggestion", "") or (f.get("why", "") or "")[:_MSG_MAX]

            rule_id = f.get("rule_id", "")
            sev_txt = sev_label + (f"\n{rule_id}" if rule_id else "")
            row_data = [str(ri), sev_txt, where, problem_txt, fix or "—"]
            for ci, val in enumerate(row_data):
                cell = prob_table.cell(ri, ci)
                cell.text = str(val)
                cell.width = col_widths[ci]
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(8)
                        if ci == 1:
                            run.bold = True
                            run.font.color.rgb = RGBColor(*_severity_rgb(sev))
                _set_cell_shading(cell, hex_c)
        # Fix header column widths too
        for ci in range(5):
            prob_table.cell(0, ci).width = col_widths[ci]

        if len(problem_list) > _MAX_PROBLEM_BLOCKS:
            doc.add_paragraph()
            _add_para(
                f"({len(problem_list) - _MAX_PROBLEM_BLOCKS} constat(s) "
                f"supplémentaires non détaillés ici — liste complète disponible "
                f"dans l'analyse en ligne.)",
                italic=True, size=8, color=(108, 117, 125)
            )

        # ── Per-problem resolution detail ──────────────────
        doc.add_paragraph()
        _add_heading("Détail et résolution de chaque problème", level=3)
        _add_para(
            "Pour chaque problème du tableau : sa localisation exacte, la "
            "nature précise du problème, la marche à suivre pour le corriger, "
            "et un exemple lorsque pertinent.",
            italic=True, size=9, color=(108, 117, 125)
        )
        doc.add_paragraph()

        for idx, (sev_label, f) in enumerate(shown, 1):
            sev = f.get("severity", "warning")
            sev_rgb = _severity_rgb(sev)
            rule_id = f.get("rule_id", "")
            sec = f.get("section", "")
            location = f.get("user_location", "")
            message = (f.get("message", "") or "")[:_MSG_MAX]
            excerpt = (f.get("user_excerpt", "") or "")[:_EXCERPT_MAX]
            why = (f.get("why", "") or "")[:_MSG_MAX]
            fix = f.get("fix_suggestion", "")
            example = _example_for(f)

            head = f"{idx}. [{sev_label}] {rule_id}" + (f" — {sec}" if sec else "")
            _add_para(head, bold=True, size=10, color=sev_rgb)

            where_parts = []
            if sec:
                where_parts.append(f"section « {sec} »")
            if location:
                where_parts.append(location)
            if where_parts:
                _add_para("Où : " + " — ".join(where_parts), size=9)
            if excerpt:
                _add_para(f"Extrait concerné : « {excerpt} »",
                          italic=True, size=8, color=(90, 90, 90))
            if message:
                _add_para(f"Problème : {message}", size=9)
            if why:
                _add_para(f"Pourquoi : {why}", size=9, color=(80, 80, 80))
            if fix:
                _add_para(f"→ Correction : {fix}", size=9, color=(0, 100, 0))
            if example:
                _add_para(f"Exemple : {example}", italic=True, size=8.5,
                          color=(0, 51, 102))
            doc.add_paragraph()
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════
    # 3. COUVERTURE DES SECTIONS
    # ═══════════════════════════════════════════════════════
    _add_heading("3. Couverture des Sections du Template", level=2)

    if sections_missing:
        _add_para(
            f"❌ {len(sections_missing)} section(s) obligatoire(s) manquante(s) :",
            bold=True, size=10, color=(220, 53, 69)
        )
        for s in sections_missing:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(str(s))
            run.font.size = Pt(9)
            run.bold = True
            run.font.color.rgb = RGBColor(220, 53, 69)
    else:
        _add_para(
            "✅ Toutes les sections obligatoires du template sont présentes.",
            bold=True, size=10, color=(40, 167, 69)
        )

    if sections_found:
        _add_para(
            f"{len(sections_found)} sections détectées dans le document.",
            size=8, color=(108, 117, 125)
        )
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════
    # 4. RECOMMANDATIONS
    # ═══════════════════════════════════════════════════════
    _add_heading("4. Recommandations", level=2)

    for rec in _generate_recommendations(report)[:_MAX_RECOMMENDATIONS]:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(rec)
        run.font.size = Pt(9.5)
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════
    # 5. PÉRIMÈTRE DE L'ANALYSE (audit trail, compact)
    # ═══════════════════════════════════════════════════════
    _add_heading("5. Périmètre de l'Analyse", level=2)

    wg_count = rules_used.get("writing_guide_rules_count", 0)
    wg_checked = rules_used.get("writing_guide_rules_checked", 0)
    mandatory_count = rules_used.get("mandatory_sections_count", 0)
    checked_ids = rules_used.get("checked_rule_ids", [])
    source_docs = rules_used.get("source_documents", [])
    extraction_ok = rules_used.get("extraction_ok", True)
    text_length = report.get("textLength", 0)

    _add_para(
        f"Document analysé en intégralité ({text_length:,} caractères extraits). "
        f"Vérifications : {mandatory_count} sections obligatoires du template + "
        f"{wg_checked}/{wg_count} règles du guide d'écriture "
        f"({counts.get('pass', 0)} vérifications conformes). "
        f"Extraction des règles : {'réussie' if extraction_ok else 'AVEC ERREURS'}.",
        size=9
    )
    if checked_ids:
        _add_para("Règles vérifiées : " + ", ".join(checked_ids),
                  size=8, color=(108, 117, 125))
    unchecked_ids = rules_used.get("unchecked_rule_ids", [])
    if unchecked_ids:
        _add_para(
            f"Règles du guide non couvertes par l'analyse automatique "
            f"({len(unchecked_ids)}) : " + ", ".join(unchecked_ids)
            + " — à vérifier manuellement.",
            size=8, color=(176, 122, 0)
        )
    if source_docs:
        _add_para("Documents de référence : " + " ; ".join(source_docs),
                  size=8, color=(108, 117, 125))
    _add_para(
        "Méthode : analyse déterministe basée sur des preuves — chaque constat "
        "cite la règle source et l'extrait exact du document (double preuve). "
        "Détail complet disponible dans l'analyse en ligne (interface LEON).",
        italic=True, size=8, color=(108, 117, 125)
    )

    extraction_errors = rules_used.get("errors", [])
    if extraction_errors:
        _add_para("Erreurs d'extraction des règles :", bold=True, size=8, color=(220, 53, 69))
        for err in extraction_errors:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(str(err))
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(220, 53, 69)

    # ── Footer (page numbers) ──────────────────────────────
    _add_page_number_footer(section)

    # ── Save to bytes ───────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ═══════════════════════════════════════════════════════════════════
# RECOMMENDATION GENERATOR
# ═══════════════════════════════════════════════════════════════════

def _generate_recommendations(report: Dict) -> List[str]:
    """Generate actionable, prioritized recommendations from the report."""
    recs: List[str] = []

    verdict = report.get("verdict", "UNKNOWN")
    counts = report.get("summaryCounts", {})
    errors = counts.get("errors", 0)
    warnings = counts.get("warnings", 0)
    scores = report.get("scores", {})
    sections_missing = report.get("sectionsMissing", [])

    # Verdict-based recommendation
    if verdict == "GOOD":
        recs.append(
            "Le document est globalement conforme. Traiter les avertissements "
            "restants pour la prochaine révision."
            if warnings else
            "Le document est conforme. Aucune action corrective nécessaire."
        )
    elif verdict == "ACCEPTABLE_WITH_FIXES":
        recs.append(
            f"Corriger les {errors} erreur(s) pour atteindre le verdict GOOD ; "
            f"traiter ensuite les {warnings} avertissement(s)."
        )
    elif verdict == "NOT_RELIABLE":
        recs.append(
            f"Révision approfondie nécessaire avant soumission "
            f"({errors} erreurs, {warnings} avertissements)."
        )
    else:  # NON_COMPLIANT
        recs.append(
            f"Réécriture significative nécessaire en suivant le template CTS "
            f"et le guide d'écriture ({errors} erreurs critiques)."
        )

    if sections_missing:
        recs.append(
            f"Ajouter les {len(sections_missing)} section(s) obligatoire(s) "
            "manquante(s) : "
            + ", ".join(str(s) for s in sections_missing[:8])
            + ("…" if len(sections_missing) > 8 else "")
            + "."
        )

    if scores.get("template_cleanliness", 1) < 0.80:
        recs.append(
            "Supprimer les placeholders restants (<<…>>, TBD, XXX) "
            f"(propreté du template : {scores.get('template_cleanliness', 0):.0%})."
        )

    if scores.get("requirements_quality", 1) < 0.60:
        recs.append(
            "Renforcer les exigences : 'shall' pour chaque exigence obligatoire, "
            "ID unique par exigence, traçabilité vers les exigences amont "
            f"(score actuel : {scores.get('requirements_quality', 0):.0%})."
        )

    if scores.get("writing_guide_compliance", 1) < 0.60:
        recs.append(
            "Consulter le guide d'écriture CTS pour les règles de rédaction "
            f"(conformité actuelle : {scores.get('writing_guide_compliance', 0):.0%})."
        )

    return recs
