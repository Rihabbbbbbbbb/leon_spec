"""
Pipeline d'ingestion des documents de référence.
Extrait le texte de manière structurée (sections, tableaux, instructions),
découpe en chunks, génère les embeddings et sauvegarde l'index.
"""
import json
import re
from pathlib import Path
from typing import List, Dict, Any, Optional
from docx import Document
from docx.oxml.ns import qn

from app.config import REFS_DIR, INDEX_PATH, DEFAULT_CHUNK_SIZE, DEFAULT_CHUNK_STRATEGY
from app.chunking import chunk_document
from app.embeddings import get_embeddings_batch


# ── Patterns for template/non-normative content detection ──────────
TEMPLATE_PLACEHOLDER_RE = re.compile(r'<<[^>]*>>')
TBD_PLACEHOLDER_RE = re.compile(r'\b(TBD|XXX|TODO)\b', re.IGNORECASE)
TABLE_OF_CONTENTS_RE = re.compile(r'^\s*(\d+(?:\.\d+)*)?\s*\.+\s*\d+\s*$')  # TOC line with page number
PAGE_NUMBER_ONLY_RE = re.compile(r'^\s*\d+\s*$')  # Standalone page numbers


def _is_template_instruction(text: str) -> bool:
    """Detect if text is a template instruction (placeholder, red text marker, etc.)."""
    if TEMPLATE_PLACEHOLDER_RE.search(text):
        return True
    if TBD_PLACEHOLDER_RE.search(text) and len(text.strip()) < 50:
        return True
    return False


def _is_toc_line(text: str) -> bool:
    """Detect Table of Contents lines (section title ... page number)."""
    return bool(TABLE_OF_CONTENTS_RE.match(text.strip()))


def _get_heading_level(para) -> Optional[int]:
    """Extract heading level from a paragraph's Word style (1-6) or None if not a heading."""
    style_name = para.style.name if para.style else ""
    # Standard Word heading styles
    if style_name.startswith("Heading") or style_name.startswith("heading"):
        try:
            return int(style_name.split()[-1])
        except ValueError:
            return 1
    # French/Stellantis heading styles (Titre 1, Titre 2, etc.)
    if style_name.startswith("Titre") or style_name.startswith("titre"):
        try:
            return int(style_name.split()[-1])
        except ValueError:
            return 1
    # Style names containing heading numbers
    for prefix in ["H", "HN", "Head"]:
        if style_name.startswith(prefix):
            try:
                return int(re.findall(r'\d+', style_name)[0])
            except (IndexError, ValueError):
                pass
    return None


def _is_red_text(para) -> bool:
    """Check if paragraph contains red text (template instructions in Stellantis docs)."""
    for run in para.runs:
        if run.font.color and run.font.color.rgb:
            rgb = str(run.font.color.rgb)
            # Red or near-red
            if rgb.startswith("FF") or rgb.startswith("FE") or rgb.startswith("C"):
                return True
    return False


def extract_docx_text(file_path: str) -> List[Dict[str, Any]]:
    """
    Extraction structurée d'un fichier .docx avec métadonnées de section,
    type de bloc, et détection des instructions template.

    Pour chaque élément on conserve :
    - source_file : nom du fichier
    - block_id   : index séquentiel
    - block_type : 'heading', 'paragraph', 'table', 'table_row', 'template_instruction'
    - heading_level : 1-6 si heading, None sinon
    - section_context : titre de la section parente la plus proche
    - text       : contenu textuel nettoyé
    - is_template : True si instruction/placeholder template

    Args:
        file_path: Chemin absolu vers le fichier .docx

    Returns:
        Liste de blocs textuels structurés
    """
    doc = Document(file_path)
    blocks: List[Dict[str, Any]] = []
    source_name = Path(file_path).name
    idx = 0
    current_section: str = ""
    section_stack: List[str] = []

    def _update_section(heading_text: str, level: int):
        """Maintain a section stack for context tracking."""
        nonlocal section_stack, current_section
        # Trim stack to parent level
        section_stack = section_stack[:level - 1] if level > 1 else []
        section_stack.append(heading_text)
        current_section = " > ".join(section_stack)

    # ── Phase 1: Extract paragraphs with structure ─────────────────
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        heading_level = _get_heading_level(para)
        is_red = _is_red_text(para)
        is_template = _is_template_instruction(text) or is_red
        is_toc = _is_toc_line(text)

        # Skip pure TOC lines (section name ... page number)
        if is_toc and not heading_level:
            continue

        # Skip standalone page numbers
        if PAGE_NUMBER_ONLY_RE.match(text):
            continue

        if heading_level:
            block_type = "heading"
            _update_section(text, heading_level)
        elif is_template and not heading_level:
            block_type = "template_instruction"
            # Still track as content but mark for filtering
        else:
            block_type = "paragraph"

        blocks.append({
            "source_file": source_name,
            "block_id": idx,
            "block_type": block_type,
            "heading_level": heading_level,
            "section_context": current_section,
            "text": text,
            "is_template": is_template,
        })
        idx += 1

    # ── Phase 2: Extract tables preserving structure ────────────────
    for table in doc.tables:
        rows_data: List[str] = []
        header_row: Optional[str] = None
        row_count = len(table.rows)
        col_count = len(table.columns) if table.columns else 0

        for row_idx, row in enumerate(table.rows):
            cell_texts = []
            for cell in row.cells:
                ct = cell.text.strip()
                if ct:
                    cell_texts.append(ct)
            if not cell_texts:
                continue

            row_text = " | ".join(cell_texts)

            # First row is typically the header
            if row_idx == 0 and col_count >= 2:
                header_row = row_text
                # Emit the header as a labeled table marker
                blocks.append({
                    "source_file": source_name,
                    "block_id": idx,
                    "block_type": "table",
                    "heading_level": None,
                    "section_context": current_section,
                    "text": f"[TABLE START: {row_text}]",
                    "is_template": False,
                })
                idx += 1
                continue

            # Detect if this is a requirement table (has columns like "Requirement Number | Description | ...")
            is_req_table = bool(
                header_row and
                re.search(r'(requirement|exigence|number|numéro|description|input|validation)',
                          header_row, re.IGNORECASE)
            )

            row_block_type = "table_row"
            if is_req_table:
                # Keep requirement rows as individual entries for retrieval
                row_block_type = "table_row"

            blocks.append({
                "source_file": source_name,
                "block_id": idx,
                "block_type": row_block_type,
                "heading_level": None,
                "section_context": current_section,
                "text": row_text,
                "is_template": _is_template_instruction(row_text),
                "table_header": header_row,
            })
            idx += 1

        # Close table marker
        if row_count > 0 and col_count >= 2:
            blocks.append({
                "source_file": source_name,
                "block_id": idx,
                "block_type": "table_end",
                "heading_level": None,
                "section_context": current_section,
                "text": "[TABLE END]",
                "is_template": False,
            })
            idx += 1

    return blocks


def extract_docx_text_flat(file_path: str) -> List[Dict[str, Any]]:
    """
    Extraction simplifiée (fallback) — paragraphes et tableaux à plat.
    Utile quand la structure n'est pas critique pour le chunking.
    """
    doc = Document(file_path)
    blocks: List[Dict[str, Any]] = []
    source_name = Path(file_path).name
    idx = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            blocks.append({
                "source_file": source_name,
                "block_id": idx,
                "text": text,
            })
            idx += 1

    for table in doc.tables:
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)
            if row_texts:
                blocks.append({
                    "source_file": source_name,
                    "block_id": idx,
                    "text": " | ".join(row_texts),
                })
                idx += 1

    return blocks


def ingest_reference_doc(
    file_path: str,
    chunk_strategy: str = DEFAULT_CHUNK_STRATEGY,
    chunk_size: int = DEFAULT_CHUNK_SIZE,
) -> List[Dict[str, Any]]:
    """
    Ingère un document de référence complet :
    extraction → chunking → embedding → stockage.

    Args:
        file_path: Chemin du fichier .docx
        chunk_strategy: Stratégie de chunking ('paragraphs' ou 'sections')
        chunk_size: Nombre de paragraphes par chunk

    Returns:
        Liste de chunks indexés avec leurs embeddings
    """
    # Étape 1 : extraction
    blocks = extract_docx_text(file_path)
    if not blocks:
        print(f"[WARN] Aucun contenu extrait de {file_path}")
        return []

    # Étape 2 : chunking
    chunks = chunk_document(blocks, strategy=chunk_strategy, chunk_size=chunk_size)
    if not chunks:
        print(f"[WARN] Aucun chunk généré pour {file_path}")
        return []

    # Étape 3 : embedding batch
    try:
        vectors = get_embeddings_batch(chunks)
    except RuntimeError as e:
        print(f"[ERROR] Échec embedding pour {file_path}: {e}")
        raise

    # Étape 4 : assemblage
    indexed = []
    source_name = Path(file_path).name
    for idx, (chunk_text, vector) in enumerate(zip(chunks, vectors)):
        indexed.append({
            "source_file": source_name,
            "chunk_id": idx,
            "text": chunk_text,
            "embedding": vector,
        })

    return indexed


def load_reference_index() -> List[Dict[str, Any]]:
    """
    Charge l'index de référence depuis le disque.

    Returns:
        Liste de chunks indexés (liste vide si l'index n'existe pas)
    """
    if not INDEX_PATH.exists():
        return []
    with open(INDEX_PATH, "r", encoding="utf-8") as f:
        return json.load(f)


def save_reference_index(index: List[Dict[str, Any]]) -> None:
    """
    Sauvegarde l'index de référence sur le disque.

    Args:
        index: Liste de chunks indexés
    """
    INDEX_PATH.parent.mkdir(parents=True, exist_ok=True)
    with open(INDEX_PATH, "w", encoding="utf-8") as f:
        json.dump(index, f, ensure_ascii=False)


def build_index(
    refs_folder: str = None,
    chunk_strategy: str = DEFAULT_CHUNK_STRATEGY,
    chunk_size: int = DEFAULT_CHUNK_SIZE,
) -> List[Dict[str, Any]]:
    """
    Construit l'index complet à partir de tous les fichiers .docx
    présents dans le dossier de références.

    Args:
        refs_folder: Chemin du dossier (défaut: config REFS_DIR)
        chunk_strategy: Stratégie de chunking
        chunk_size: Taille des chunks

    Returns:
        Liste complète de chunks indexés
    """
    folder = Path(refs_folder) if refs_folder else REFS_DIR

    if not folder.exists():
        raise FileNotFoundError(f"Dossier de références introuvable : {folder}")

    docx_files = list(folder.glob("*.docx"))
    if not docx_files:
        raise FileNotFoundError(f"Aucun fichier .docx trouvé dans {folder}")

    all_items: List[Dict[str, Any]] = []
    for file_path in docx_files:
        print(f"[INFO] Ingestion de {file_path.name}...")
        try:
            items = ingest_reference_doc(
                str(file_path),
                chunk_strategy=chunk_strategy,
                chunk_size=chunk_size,
            )
            all_items.extend(items)
            print(f"  -> {len(items)} chunks generes")
        except Exception as e:
            print(f"  -> [ERROR] {e}")

    # Sauvegarde
    save_reference_index(all_items)
    print(f"[INFO] Index sauvegardé : {len(all_items)} chunks au total")
    return all_items


# Point d'entrée standalone
if __name__ == "__main__":
    build_index()
