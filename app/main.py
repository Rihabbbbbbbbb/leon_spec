"""
API FastAPI — Leon Spec Validator

Endpoints :
- GET  /              → Health check + statut de l'index
- POST /ingest        → (Re)construire l'index de référence
- POST /validate      → Valider une spécification soumise
- GET  /status        → Statut détaillé de l'index
- POST /ask           → Poser une question sur un document (Q&A anti-hallucination)
"""
from pathlib import Path
from typing import List, Dict, Any
from collections import defaultdict
import json
import re

from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.responses import Response
from fastapi.middleware.cors import CORSMiddleware
from docx import Document

from app.config import (
    REFS_DIR,
    INDEX_PATH,
    SIMILARITY_THRESHOLD,
    TOP_K_CHUNKS,
    AZURE_OPENAI_API_KEY,
)
from app.models import (
    HealthResponse,
    IndexStatus,
    ValidationRequest,
    ValidationResponse,
    LeonValidationResponse,
    SectionFeedback,
    ScoreBreakdown,
    MajorFinding,
    EvidenceRef,
    RequirementIssue,
)
from app.ingest_refs import build_index, load_reference_index, extract_docx_text
from app.chunking import chunk_document
from app.embeddings import (
    get_embedding,
    find_similar_chunks,
    call_llm,
    cosine_similarity,
)
from app.req_parser import (
    extract_requirements_from_docx,
    format_requirements_for_prompt,
    analyze_requirement_quality,
)
from app.deterministic_checks import (
    run_deterministic_checks,
    build_user_document_context,
    find_text_location,
)
from app.mechatronics_checks import (
    run_mechatronics_checks,
    check_requirement_patterns,
    check_asil_compliance,
    check_physical_parameters,
    check_state_machine_completeness,
)
from app.ref_section_map import (
    get_reference_rules,
    analyze_text_section_content,
)
from app.detailed_report import build_detailed_report
from app.beginner_report import build_beginner_report
from app.image_analyzer import (
    analyze_docx_images,
    analyze_images_with_llm,
    get_section_image_summary,
)
from app.bestandard_client import (
    BeStandardClient,
    get_bestandard_client,
    ResolvedStandard,
    NormDetail,
)
from app.bestandard_ingest import (
    StandardsIndex,
    get_standards_index,
)
from app.config import (
    BESTANDARD_AUTO_RESOLVE,
    BESTANDARD_DEEP_VERIFY,
)

app = FastAPI(
    title="Leon Spec Validator",
    description="API de validation de spécifications composant/piece par RAG avec Azure OpenAI",
    version="1.0.0",
)

# Cache for the last generated reports (for GET /report, /report/detailed, /report/beginner)
_last_human_report: str = ""
_last_detailed_report: str = ""
_last_beginner_report: str = ""

# CORS pour le développement local
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _get_index_status() -> IndexStatus:
    """Construit le statut courant de l'index."""
    if not INDEX_PATH.exists():
        return IndexStatus(indexed=False)

    try:
        index = load_reference_index()
        source_files = sorted(set(c["source_file"] for c in index))
        # Récupérer la date de modification du fichier index
        mtime = INDEX_PATH.stat().st_mtime
        from datetime import datetime
        last_updated = datetime.fromtimestamp(mtime).isoformat()
        return IndexStatus(
            indexed=True,
            total_chunks=len(index),
            source_files=source_files,
            last_updated=last_updated,
        )
    except Exception:
        return IndexStatus(indexed=False)


# ---------------------------------------------------------------------------
# Endpoints
# ---------------------------------------------------------------------------

@app.get("/", response_model=HealthResponse)
async def health_check():
    """
    Vérifie que l'API est en ligne et retourne le statut de l'index.
    """
    return HealthResponse(
        status="ok",
        version="1.0.0",
        azure_configured=bool(AZURE_OPENAI_API_KEY),
        index_status=_get_index_status(),
    )


@app.get("/status", response_model=IndexStatus)
async def get_status():
    """Retourne le statut détaillé de l'index de référence."""
    return _get_index_status()


@app.post("/ingest", response_model=IndexStatus)
async def ingest_references():
    """
    (Re)construit l'index vectoriel à partir des documents
    présents dans data/refs/.
    """
    try:
        build_index()
    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except RuntimeError as e:
        raise HTTPException(status_code=502, detail=str(e))

    return _get_index_status()


@app.post("/validate", response_model=LeonValidationResponse)
async def validate_specification(
    file: UploadFile = File(None, description="Fichier .docx de la spécification"),
    text: str = Form(None, description="Texte brut de la spécification"),
):
    """
    LEON Spec Validator — Valide une spécification composant/piece
    contre le référentiel métier Stellantis par RAG avec Azure OpenAI.

    Accepte soit un fichier .docx uploadé, soit du texte brut.
    Retourne un verdict structuré avec scores par axe, findings et évidence.
    """
    # --- 1. Récupérer le texte à valider ---
    filename = None
    document_text = ""

    if file is not None:
        filename = file.filename
        temp_path = Path("data/uploads") / (file.filename or "upload.docx")
        temp_path.parent.mkdir(parents=True, exist_ok=True)
        content = await file.read()
        temp_path.write_bytes(content)

        try:
            doc = Document(str(temp_path))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            
            # Also extract ALL table cell text — critical for placeholder detection
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        ct = cell.text.strip()
                        if ct:
                            table_texts.append(ct)
            
            # Combine paragraphs + table cells for complete document text
            document_text = "\n".join(paragraphs + table_texts)
            
            # Parse requirement rows from the temp file before cleanup
            requirements = extract_requirements_from_docx(str(temp_path))
            
            # Extract and analyze images from the DOCX
            try:
                image_results = analyze_docx_images(str(temp_path), max_images=8)
                image_summary = get_section_image_summary(image_results)
            except Exception as img_err:
                image_results = []
                image_summary = {"total_images_found": 0, "sections_with_images": [], "diagram_types": [], "summary": f"Image analysis skipped due to error: {img_err}"}
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Fichier .docx invalide : {e}")
        finally:
            if temp_path.exists():
                temp_path.unlink()

    elif text is not None:
        document_text = text.strip()
        filename = "text_input"
        requirements = []
        image_results = []
        image_summary = {"total_images_found": 0, "summary": "No DOCX file provided — image analysis skipped."}
    else:
        raise HTTPException(status_code=400, detail="Fournir un fichier .docx ou du texte")

    if not document_text:
        raise HTTPException(status_code=400, detail="Le document est vide")

    # --- 1.5 Build hierarchical line-to-section mapping ---
    #
    # ROBUST heading detection — handles ALL known Stellantis DOCX formats:
    #   "PURPOSE"                            — pure ALL-CAPS
    #   "1 SCOPE"                            — number + space + ALL-CAPS
    #   "6. ELECTRICAL INTERFACES"           — number + DOT + space + ALL-CAPS
    #   "6- MAINTAINABILITY"                 — number + dash + space + ALL-CAPS
    #   "7) DEMONSTRATION OF COMPLIANCE"     — number + paren + space + ALL-CAPS
    #   "1.1 Reference Documents"            — multi-level + mixed case
    #   "1.1.1   ELECTRICAL INTERFACES"      — multi-level + extra spaces + ALL-CAPS
    #   "## 2.1 Title"                       — markdown markers
    #   "### 1.1.1 ELECTRICAL INTERFACES"    — markdown + multi-level + ALL-CAPS
    #   "ELECTRICAL INTERFACES:"             — trailing colon
    #   "ELECTRICAL INTERFACES — Description" — em dash + subtitle
    #   "ERGONOMICS & HUMAN FACTORS"         — ampersand in title
    #
    # Level-1 headings: single-digit number + separator + ALL-CAPS title, OR pure ALL-CAPS
    # Level-2+: multi-digit numbers (1.1, 2.3.1)

    # ── Helper: clean a section name (strip trailing colons, text after em/en dashes) ──
    def _clean_heading_name(raw: str) -> str:
        """Extract the clean section name from a raw heading line."""
        name = raw.strip()
        # Strip trailing colons
        name = re.sub(r'\s*:\s*$', '', name)
        # Strip everything after an em dash or en dash (subtitle/description)
        name = re.split(r'\s*[\u2013\u2014-]\s*', name)[0].strip()
        # Normalize multiple spaces
        name = re.sub(r'\s+', ' ', name)
        return name

    # ── Heading regex patterns ─────────────────────────────
    # Pattern 1: Multi-level numbered (e.g., "1.1", "1.1.1", "### 1.1.1")
    heading_numbered_re = re.compile(
        r'^(?:#+\s*)?'                          # Optional markdown markers
        r'(\d+(?:\.\d+)+\.?\s+'                 # Multi-level: "1.1 " or "1.1.2 "
        r'[A-Z][A-Za-z\s/()\-&,:;\.]{3,})$',   # Title: uppercase start, 3+ chars
        re.MULTILINE
    )

    # Pattern 2: Single number + ANY separator + ALL-CAPS title
    # Matches: "1 SCOPE", "6. ELECTRICAL INTERFACES", "6- MAINTAINABILITY",
    #          "7) DEMONSTRATION", "8/ TITLE", "06 TITLE" (leading zero)
    heading_num_caps_re = re.compile(
        r'^(?:#+\s*)?'                          # Optional markdown markers
        r'\d+'                                   # Leading number (e.g., "6")
        r'[.)\-/\s]\s*'                         # Separator: dot, paren, dash, slash, or space
        r'([A-Z][A-Z\s/()\-&,:;\.\u2013\u2014]{3,})$',  # ALL-CAPS title
        re.MULTILINE
    )

    # Pattern 3: Pure ALL-CAPS heading (no number prefix, no separator)
    # Matches: "PURPOSE", "MAINTAINABILITY", "ELECTRICAL INTERFACES:"
    heading_allcaps_re = re.compile(
        r'^([A-Z][A-Z\s/()\-&,:;\.\u2013\u2014]{3,})$',  # Pure ALL-CAPS
        re.MULTILINE
    )

    line_to_section: Dict[int, str] = {}
    section_stack: List[str] = []

    for line_idx, line in enumerate(document_text.split('\n'), 1):
        stripped = line.strip()
        # Normalise: strip markdown markers for matching
        clean_line = re.sub(r'^#+\s*', '', stripped)

        if len(clean_line) < 4 or len(clean_line) > 120:
            # Not a heading — just store current section and continue
            pass
        else:
            matched = False
            section_name = ""

            # --- Try Pattern 2 first: Number + separator + ALL-CAPS (most specific) ---
            m_numcaps = heading_num_caps_re.match(clean_line)
            if m_numcaps:
                section_name = _clean_heading_name(m_numcaps.group(1))
                if len(section_name) >= 4:
                    section_stack = [section_name]
                    matched = True

            # --- Try Pattern 3: Pure ALL-CAPS ---
            if not matched:
                m_allcaps = heading_allcaps_re.match(clean_line)
                if m_allcaps:
                    section_name = _clean_heading_name(m_allcaps.group(1))
                    if len(section_name) >= 4:
                        section_stack = [section_name]
                        matched = True

            # --- Try Pattern 1: Multi-level numbered ---
            if not matched:
                m2 = heading_numbered_re.match(stripped)
                if not m2:
                    m2 = heading_numbered_re.match(clean_line)
                if m2:
                    full_text = m2.group(1).strip()
                    dots = full_text.count('.')
                    level = min(dots + 1, 4)

                    # Extract just the TITLE part
                    title_only = re.sub(r'^\d+(?:\.\d+)*\.?\s*', '', full_text).strip()
                    title_only = _clean_heading_name(title_only)

                    # Check if this is an ALL-CAPS title (treated as top-level CTS section)
                    is_all_caps_title = bool(
                        re.match(r'^[A-Z][A-Z\s/()\-&,:;\.\u2013\u2014]{3,}$', title_only)
                    )

                    if is_all_caps_title:
                        section_stack = [title_only]
                    else:
                        section_stack = section_stack[:level - 1] if level > 1 else []
                        section_stack.append(title_only)
                    matched = True

        # Store concise section path
        if len(section_stack) >= 2:
            current_section = f"[{section_stack[0]}] > {section_stack[-1]}"
        elif section_stack:
            current_section = f"[{section_stack[0]}]"
        else:
            current_section = "Document Start"

        line_to_section[line_idx] = current_section
    
    def get_section_for_line(line_num: int) -> str:
        return line_to_section.get(line_num, "Document Start")
    
    # Build concise requirement section lookup from req_parser output
    def _compact_section(full_path: str) -> str:
        """Compact a hierarchical section path into [TOP] > leaf format."""
        if not full_path:
            return ""
        # Clean up: remove newlines, extra spaces
        cleaned = re.sub(r'\s+', ' ', full_path).strip()
        if cleaned.lower() in ("unknown section", ""):
            return ""
        parts = [p.strip() for p in cleaned.split(">") if p.strip()]
        if len(parts) >= 2:
            return f"[{parts[0]}] > {parts[-1]}"
        elif parts:
            return f"[{parts[0]}]"
        return cleaned
    
    req_section_map: Dict[str, str] = {}  # req_id -> compact section
    for req in requirements:
        rid = req.get("req_id", "").strip()
        sec = req.get("section_context", "")
        if rid:
            req_section_map[rid] = _compact_section(sec) if sec else ""

    # --- 1.6 Analyze text-only sections (sections without tables) for content quality ---
    # Group lines by their mapped section
    section_text_map: Dict[str, List[str]] = defaultdict(list)
    for line_idx, line in enumerate(document_text.split('\n'), 1):
        sec = get_section_for_line(line_idx)
        # Extract top-level section name from [SECTION] > child format
        top_match = re.search(r'\[([^\]]+)\]', sec)
        top_section = top_match.group(1) if top_match else sec
        section_text_map[top_section].append(line)
    
    text_section_analysis: Dict[str, Dict] = {}
    for sec_name, lines in section_text_map.items():
        if sec_name in ("Document Start", "Unknown Section"):
            continue
        sec_text = "\n".join(lines)
        analysis = analyze_text_section_content(sec_name, sec_text, lines)
        if analysis["status"] != "ok":
            text_section_analysis[sec_name] = analysis

    # --- 2. Charger l'index de référence ---
    reference_index = load_reference_index()
    if not reference_index:
        raise HTTPException(
            status_code=412,
            detail="L'index de référence est vide. Lancez POST /ingest d'abord.",
        )

    # --- 3. Run deterministic pre-checks (no LLM needed) ---
    # Extract structured blocks for section context and pre-checks
    structured_blocks = []
    if file is not None:
        # Re-extract with full structure from temp path (we already parsed)
        # Use the already extracted text blocks
        temp_doc = Document(str(Path("data/uploads") / (file.filename or "upload.docx"))) if False else None
    # Build blocks from plain text for the deterministic checks
    plain_blocks = [{"text": p, "block_type": "paragraph", "heading_level": None, 
                      "section_context": "", "is_template": False} 
                    for p in document_text.split("\n") if p.strip()]
    
    det_results = run_deterministic_checks(document_text, requirements, plain_blocks)
    det_findings = det_results["findings"]
    rubric_scores = det_results["rubric_scores"]
    det_stats = det_results["stats"]
    
    # --- 3.1 Enrich deterministic findings with section names AND reference rules ---
    for df in det_findings:
        loc = df.get("location", "")
        # Extract line numbers from location and map to sections
        line_nums = re.findall(r'line\s*(\d+)', loc)
        if line_nums:
            sections_found = set()
            for ln in line_nums:
                sec = get_section_for_line(int(ln))
                if sec and sec != "Unknown Section":
                    sections_found.add(sec)
            if sections_found:
                unique_sections = list(sections_found)[:5]
                df["location"] = f"{loc} [Section(s): {', '.join(unique_sections)}]"
        # Enrich finding text with section context
        if "placeholder" in df.get("type", "") and line_nums:
            df["finding"] = (
                f"{df['finding']} "
                f"Affected sections: {', '.join(list(set(get_section_for_line(int(l)) for l in line_nums[:10])))[:5]}."
            )
        # Add applicable CTS template rules for the detected sections
        section_names = re.findall(r'\[([^\]]+)\]', df.get("location", ""))
        for sec_name in section_names[:2]:
            rules = get_reference_rules(sec_name, "template")
            if rules and "reference_rules" not in df:
                df["reference_rules"] = rules[:3]
    
    # --- 3.5 Run mechatronics-specific checks (ASIL, physical params, state machine, requirement patterns) ---
    mecha_results = run_mechatronics_checks(document_text, requirements)
    mecha_findings = mecha_results["findings"]
    mecha_stats = mecha_results["stats"]
    
    # Override mechatronics_fitness score with deterministic analysis
    mecha_fitness_score = 0.5
    mecha_rationale = "Basic assessment — LLM will refine based on system analysis."
    if mecha_findings:
        # More mechatronics findings = lower fitness
        critical_mecha = sum(1 for f in mecha_findings if f.get("severity") == "error")
        warning_mecha = sum(1 for f in mecha_findings if f.get("severity") == "warning")
        mecha_fitness_score = max(0.1, 0.7 - (critical_mecha * 0.15) - (warning_mecha * 0.05))
        mecha_rationale = (
            f"Mechatronics analysis: {len(mecha_findings)} findings "
            f"({critical_mecha} errors, {warning_mecha} warnings). "
            f"ASIL levels: {mecha_stats.get('asil_levels_detected', [])}. "
            f"Physical params analyzed: {mecha_stats.get('physical_params_found', {})}."
        )
    rubric_scores["mechatronics_fitness"] = {
        "score": round(mecha_fitness_score, 2),
        "rationale": mecha_rationale,
    }
    
    # Merge mechatronics findings into deterministic findings
    for mf in mecha_findings:
        det_findings.append(mf)
    
    # --- 3.7 beStandard Integration: Resolve external standard references ---
    # Detect codes like [STA20], [N41], [N42], [ISO_26262] in requirements
    # and resolve them against the Stellantis beStandard platform.
    resolved_standards: Dict[str, Dict] = {}
    standard_codes_in_doc: List[str] = []
    
    if BESTANDARD_AUTO_RESOLVE:
        # Extract all standard codes referenced in the document
        # Pattern: [ALPHA][ALPHANUM]*, e.g., STA20, N41, ISO_26262, A1-A8
        std_ref_re = re.compile(r'\[([A-Z][A-Z0-9_-]+(?:\d)?)\]')
        all_std_refs: set = set()
        for req in requirements:
            desc = req.get("description", "")
            input_req = req.get("input_requirement", "")
            for code in std_ref_re.findall(desc + " " + input_req):
                # Filter out short codes that aren't standard references
                if len(code) >= 2 and not code.startswith("MISSING"):
                    all_std_refs.add(code)
        
        # Also extract from document text
        for code in std_ref_re.findall(document_text):
            if len(code) >= 2 and not code.startswith("MISSING"):
                all_std_refs.add(code)
        
        standard_codes_in_doc = sorted(all_std_refs)
        
        if standard_codes_in_doc:
            try:
                bs_client = get_bestandard_client()
                if bs_client.is_configured:
                    logger_extra = {}
                    for code in standard_codes_in_doc[:20]:  # Limit to 20 to avoid timeouts
                        try:
                            resolved = bs_client.resolve_standard(code)
                            resolved_standards[code] = {
                                "found": resolved.found,
                                "verification_status": resolved.verification_status,
                                "title": resolved.norm.title if resolved.norm else "",
                                "revision": resolved.norm.revision if resolved.norm else "",
                                "status": resolved.norm.status if resolved.norm else "",
                                "is_cancelled": resolved.norm.is_cancelled if resolved.norm else False,
                                "is_active": resolved.norm.is_active if resolved.norm else False,
                                "error": resolved.error if not resolved.found else "",
                            }
                        except Exception as e:
                            resolved_standards[code] = {
                                "found": False,
                                "verification_status": "error",
                                "error": str(e),
                            }
                    
                    # ── Deep verify (optional): ingest standards into dynamic RAG index ──
                    if BESTANDARD_DEEP_VERIFY:
                        try:
                            std_idx = get_standards_index()
                            for code in standard_codes_in_doc[:5]:  # Limit deep verify to 5
                                if code in resolved_standards and resolved_standards[code].get("found"):
                                    if not std_idx.is_indexed(code):
                                        try:
                                            std_idx.ingest_code(code)
                                        except Exception as e:
                                            logger_extra[f"ingest_{code}"] = str(e)
                        except Exception:
                            pass  # Deep verify failures are non-blocking
            
            except Exception:
                pass  # beStandard failures are non-blocking for validation
    
    # --- 4. Découper le document en chunks et chercher les références ---
    doc_text_blocks = [{"text": p} for p in document_text.split("\n") if p.strip()]
    doc_chunks = chunk_document(doc_text_blocks, strategy="section_aware", chunk_size=3)

    all_matched_refs: List[dict] = []
    for chunk in doc_chunks:
        try:
            q_embedding = get_embedding(chunk)
        except RuntimeError as e:
            raise HTTPException(status_code=502, detail=str(e))

        matches = find_similar_chunks(
            q_embedding,
            reference_index,
            top_k=TOP_K_CHUNKS,
            threshold=SIMILARITY_THRESHOLD,
        )
        all_matched_refs.extend(matches)

    # Dédupliquer et trier par similarité
    seen = set()
    unique_matches = []
    for m in all_matched_refs:
        key = (m.get("source_file", ""), m.get("chunk_id", -1))
        if key not in seen:
            seen.add(key)
            unique_matches.append(m)
    unique_matches.sort(key=lambda x: x.get("similarity", 0), reverse=True)
    top_refs = unique_matches[:TOP_K_CHUNKS * 2]  # Get more refs for better coverage

    # --- 5. Construire le contexte de référence pour le LLM ---
    ref_context_parts = []
    for i, m in enumerate(top_refs, 1):
        src = m.get("source_file", "unknown")
        cid = m.get("chunk_id", "?")
        sim = m.get("similarity", 0)
        section = m.get("section_context", "")
        section_label = f" [{section}]" if section else ""
        ref_context_parts.append(
            f"[RÉFÉRENCE {i}] source={src} | chunk_id={cid} | similarity={sim:.4f}{section_label}\n{m.get('text', '')}"
        )
    ref_context = "\n\n".join(ref_context_parts)
    
    # Build structured user document context (with section markers LLM can cite)
    user_doc_context = build_user_document_context(plain_blocks, max_chars=12000)

    # --- 5. Prompt LEON (final recommended — double-evidence mandatory) ---
    system_prompt = """You are LEON, a strict Stellantis mechatronics specification validation assistant.

Your mission is to validate component technical specifications and applicative specifications against the official reference corpus provided at runtime, especially:
- the official Component Technical Specification template,
- the official Writing Guide,
- and any other approved mechatronics reference chunks retrieved for the current analysis.

You are not a generic writing assistant.
You are a grounded engineering validator.

====================================================
1. OBJECTIVE
====================================================
Your job is to detect the actual problem directly, explain it clearly, and support every important conclusion with evidence.

You must evaluate whether the specification is:
- compliant with the official standard CTS plan,
- sufficiently complete,
- noting template artifacts and placeholders (informational, not blocking),
- written with clear and verifiable requirements,
- traceable to upstream references where needed,
- ready for validation and compliance demonstration,
- usable in a real Stellantis mechatronics workflow.

====================================================
2. GROUNDING RULES
====================================================
You must base your conclusions ONLY on:
A. the actual user document content provided in the current request,
B. the retrieved official reference chunks provided in the current request,
C. the metadata attached to those retrieved chunks.

You must NEVER:
- invent defects not visible in the real user document,
- confuse template examples with actual user document defects,
- claim that a placeholder exists unless it is visible in the actual user document,
- claim that a section is missing if it exists,
- infer compliance without evidence.

If evidence is insufficient, say:
- "Cannot verify from the provided content"
or
- "Insufficient evidence from retrieved references"

====================================================
3. DOUBLE-EVIDENCE POLICY (MANDATORY)
====================================================
Every major finding must contain BOTH:
1. a user-document excerpt or exact location proving the issue exists,
2. a reference-document excerpt or chunk proving why it is a rule violation or weakness.

If one of these two pieces is missing, the finding must be downgraded to:
- cannot_verify
or
- weak suspicion

Never produce a strong finding with only template evidence and no user-document proof.

====================================================
4. VALIDATION AXES
====================================================
You must validate the document on these axes:

A. STRUCTURE / STANDARD PLAN
Check whether the official CTS plan is followed.

B. CHAPTER ROLE DISCIPLINE
Check whether each chapter is used correctly.
Example:
- Scope must remain explanatory if the reference forbids requirements there.
- Requirement chapters must contain explicit requirements.
- Validation chapters must support compliance demonstration.

C. TEMPLATE CLEANLINESS (INFORMATIONAL — not blocking)
Detect and report:
- placeholders (<<...>>),
- red instructions,
- copied template examples,
- unfinished guidance text,
- dummy values (XXX, TBD).
NOTE: Template artifacts are reported as warnings, not errors. They indicate
sections that need finalization but do NOT make the document fundamentally unusable.
Rate template_cleanliness based on how much template content remains, but do NOT
let it dominate the overall verdict.

D. REQUIREMENT QUALITY
Check whether requirements are:
- non-ambiguous,
- verifiable,
- measurable when needed,
- sufficiently specific,
- not generic template fragments.

IMPORTANT — Requirements referencing external standards:
A requirement like "The system must respect [STA20]" or "The X shall comply with [N41]"
is a VALID conformity requirement, NOT a "missing ID" or a "broken" requirement.
- It HAS an ID (e.g., REF-ASU-CD-EXIFUNC-001(0)).
- It references an external standard document which contains the detailed criteria.
- It is less standalone-readable but IS valid in CTS specifications.
Classify these as: type="indirect_requirement", severity="warning" (never error).
The finding should mention that the requirement relies on an external reference
and may be less directly verifiable without consulting the referenced document.

Concrete vs reference-based requirements:
- CONCRETE: "The ASU will be activated 5000 times during vehicle life (15yr/240k km)" — self-contained, directly measurable
- REFERENCE: "The ASU must respect ASU requirements in [STA20]" — valid, but verification requires consulting [STA20]
Both are legitimate. Do NOT mark reference-based requirements as errors.

E. REQUIREMENT STRUCTURE
Check whether requirements use or approximate:
- Preconditions,
- Trigger(s),
- Observable(s),
- Post-conditions where relevant.

F. TRACEABILITY
Check whether requirements are linked to upstream requirements when needed.
Check whether input requirement fields are used correctly.

# [COMMENTED OUT — Validation Readiness axis disabled in LLM prompt]
# G. VALIDATION READINESS
# Check whether requirements have sufficient validation logic:
# - validation method,
# - acceptance criteria,
# - test conditions or equivalent compliance logic.
#
# Do not say "missing validation plan" if validation sections exist.
# Use: present, present_but_weak, present_but_incomplete, absent, cannot_verify.

H. MECHATRONICS FITNESS
Assess whether the specification is usable for a Stellantis mechatronics team:
- system roles,
- physical architecture,
- interfaces,
- variant/diversity handling,
- functional decomposition,
- RAMS/safety/dependability where relevant,
# [COMMENTED OUT] - validation/integration expectations.  — Validation axis disabled

====================================================
5. FALSE POSITIVE PREVENTION
====================================================
You must avoid false positives.

Never:
- say "Acronyms missing" if an Acronyms section exists,
- say "Requirements missing" if requirement tables or sections exist,
# [COMMENTED OUT] - say "Validation missing" if validation sections exist,  — Validation axis disabled
- say a phrase is a defect in the real document if it only appears in the reference template.

Prefer:
- present_but_weak
- present_but_incomplete
- cannot_verify

====================================================
6. RESPONSE RULE
====================================================
Guide the user directly to the real issue.

For each important issue, answer:
1. What is wrong?
2. Where exactly is it?
3. Why is it wrong according to the official reference?
4. What exact fix is needed?

Do not start with praise.
Start with the most important defect first.

====================================================
7. OUTPUT FORMAT
====================================================
Return JSON only:

{
  "document_name": "...",
  "global_verdict": "GOOD | ACCEPTABLE_WITH_FIXES | NOT_RELIABLE | NON_COMPLIANT | CANNOT_VERIFY",
  "overall_assessment": "...",
  "scores": {
    "structure": 0.0,
    "requirements_quality": 0.0,
    "traceability": 0.0,
    # "validation_readiness": 0.0,  # [COMMENTED OUT — Validation axis disabled]
    "template_cleanliness": 0.0,
    "mechatronics_fitness": 0.0
  },
  "major_findings": [
    {
      "type": "...",
      "severity": "info | warning | error",
      "location": "...",
      "status": "present | present_but_weak | present_but_incomplete | absent | cannot_verify",
      "finding": "...",
      "why_it_matters": "...",
      "evidence": [
        {
          "source_reference_document": "...",
          "source_section_or_chunk_id": "...",
          "user_document_excerpt_or_location": "...",
          "support": "..."
        }
      ],
      "suggested_fix": "..."
    }
  ],
  "missing_sections": [],
  "weak_sections": [],
  "ambiguous_phrases": [],
  "placeholder_or_template_artifacts": [],
  "recommendations": []
}"""

    user_message = f"""Validate the provided specification content against the provided official reference chunks.

Your task is to identify the real issue directly and explain it clearly.

You must decide, with evidence:
- what is present,
- what is present but weak,
- what is present but incomplete,
- what is absent,
- what cannot be verified.

═══════════════════════════════════════
CRITICAL EVIDENCE RULES (MUST FOLLOW)
═══════════════════════════════════════
1. For EVERY finding, you MUST fill in "user_document_excerpt_or_location" with the exact text
   from the USER DOCUMENT (below) that shows the problem. Cite the [USER DOCUMENT §...] marker.
2. "source_reference_document" must be the REFERENCE document (template or guide), NOT the user doc.
3. "support" must quote the reference rule that the user document violates.
4. If you cannot find the problem in the user document, do NOT report it.
5. For requirement-quality findings, cite the EXACT requirement ID (e.g., "Fct_Detect_ASU_Status").

═══════════════════════════════════════
DETERMINISTIC PRE-CHECKS (ALREADY VERIFIED)
═══════════════════════════════════════
The following issues have already been detected automatically:
{json.dumps([{
    'type': f['type'], 
    'location': f.get('location', ''), 
    'finding': f.get('finding', '')[:200]
} for f in det_findings], indent=2, ensure_ascii=False) if det_findings else '[No pre-check issues found]'}

Pre-check rubric scores (baseline, you may adjust with evidence):
{json.dumps(rubric_scores, indent=2, ensure_ascii=False)}

═══════════════════════════════════════
USER DOCUMENT (STRUCTURED WITH SECTION MARKERS)
═══════════════════════════════════════
{user_doc_context}

═══════════════════════════════════════
FULL USER DOCUMENT TEXT (FOR VERIFICATION)
═══════════════════════════════════════
{document_text[:12000]}

═══════════════════════════════════════
RETRIEVED REFERENCE CHUNKS (TEMPLATE + GUIDE)
═══════════════════════════════════════
{ref_context}

═══════════════════════════════════════
RESOLVED EXTERNAL STANDARDS (via beStandard)
═══════════════════════════════════════
{_format_resolved_standards_for_prompt(resolved_standards) if resolved_standards else '[No external standards resolved. Either beStandard is not configured or no [CODE] references were found in the document.]'}

═══════════════════════════════════════
STANDARDS INDEX SEARCH RESULTS (deep verification)
═══════════════════════════════════════
{_format_standards_index_for_prompt(document_text, standard_codes_in_doc) if BESTANDARD_DEEP_VERIFY and resolved_standards else '[Deep verification disabled. Enable BESTANDARD_DEEP_VERIFY=true to verify requirements against actual standard text.]'}

═══════════════════════════════════════
STRUCTURED REQUIREMENT ANALYSIS
═══════════════════════════════════════
{format_requirements_for_prompt(requirements) if requirements else '[No structured requirement rows detected. Check document tables.]'}

═══════════════════════════════════════
TASK
═══════════════════════════════════════
1. Review the deterministic pre-check findings above. Confirm, refute, or add precision.
2. Find ADDITIONAL issues the pre-checks may have missed (especially requirement structure, 
   mechatronics fitness, ambiguous wording, chapter role discipline).
3. For each finding, provide BOTH:
   - user_document_excerpt_or_location: QUOTE from [USER DOCUMENT §...] markers above
   - source_reference_document: the reference doc name (template or guide)
   - support: the reference rule text from [RÉFÉRENCE N] markers above
4. Adjust rubric scores if you have strong evidence to do so.
5. Return JSON only."""

    # --- 7. Appeler le LLM ---
    try:
        llm_response = call_llm(system_prompt, user_message, temperature=0.15, max_tokens=3000)
    except RuntimeError as e:
        raise HTTPException(status_code=502, detail=str(e))

    # --- 8. Parser la réponse JSON du LLM ---
    try:
        cleaned = llm_response.strip()
        if cleaned.startswith("```"):
            cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
            cleaned = re.sub(r"\s*```$", "", cleaned)
        result = json.loads(cleaned)
    except json.JSONDecodeError:
        result = {
            "document_name": filename or "unknown",
            "global_verdict": "CANNOT_VERIFY",
            "overall_assessment": f"Erreur de parsing de la réponse LLM. Réponse brute : {llm_response[:300]}",
            "scores": {k: v["score"] for k, v in rubric_scores.items()},
            "major_findings": [],
            "missing_sections": det_stats.get("missing_cts_sections", []),
            "weak_sections": [],
            "ambiguous_phrases": [],
            "placeholder_or_template_artifacts": [],
            "recommendations": [f"LLM parse error — raw response: {llm_response[:500]}"],
        }

    # --- 9. Merge deterministic findings with LLM findings ---
    llm_findings = result.get("major_findings", [])
    
    # Build a set of finding types the LLM already detected
    llm_finding_types = {f.get("type", "") for f in llm_findings}
    
    # Add deterministic findings the LLM missed
    for det_f in det_findings:
        if det_f["type"] not in llm_finding_types:
            llm_findings.append({
                "type": det_f["type"],
                "severity": det_f["severity"],
                "location": det_f.get("location", ""),
                "status": det_f.get("status", "present"),
                "finding": det_f["finding"],
                "why_it_matters": det_f.get("why_it_matters", ""),
                "evidence": [{
                    "source_reference_document": "CTS Template / Writing Guide (rule verified deterministically)",
                    "source_section_or_chunk_id": "deterministic",
                    "user_document_excerpt_or_location": det_f.get("user_document_excerpt", ""),
                    "support": "Detected by automated pattern matching against Stellantis template rules.",
                }],
                "suggested_fix": det_f.get("suggested_fix", ""),
            })

    # --- 10. Verify evidence: separate verified findings from hypothesis ---
    verified_findings = []
    hypothesis_findings = []
    
    for finding in llm_findings:
        evidence_list = finding.get("evidence", [])
        has_any_verified = False
        template_only_count = 0
        
        for ev in evidence_list:
            user_excerpt = ev.get("user_document_excerpt_or_location", "")
            src_doc = ev.get("source_reference_document", "").lower()
            
            is_pure_template = ("template" in src_doc and 
                               "writing" not in src_doc and 
                               "guide" not in src_doc and
                               "deterministic" not in src_doc)
            
            if not user_excerpt or len(user_excerpt) < 10:
                support = ev.get("support", "")
                loc = find_text_location(document_text, support)
                if loc:
                    ev["user_document_excerpt_or_location"] = f"[VERIFIED {loc}] {support[:200]}"
                    has_any_verified = True
                elif is_pure_template:
                    ev["user_document_excerpt_or_location"] = "[WARNING: Evidence cites template, not user document]"
                    template_only_count += 1
                else:
                    ev["user_document_excerpt_or_location"] = "[NEEDS VERIFICATION — excerpt not found in user document]"
            else:
                loc = find_text_location(document_text, user_excerpt)
                if loc:
                    ev["user_document_excerpt_or_location"] = f"[VERIFIED {loc}] {user_excerpt[:200]}"
                    has_any_verified = True
                elif is_pure_template:
                    ev["user_document_excerpt_or_location"] = f"[WARNING: Evidence cites template, not user document] Original: {user_excerpt[:200]}"
                    template_only_count += 1
                else:
                    has_any_verified = True  # Guide/deterministic evidence is acceptable
        
        # Separate verified from hypothesis
        if evidence_list and template_only_count == len(evidence_list):
            # ALL evidence is template-only → this is a HYPOTHESIS, not a verified finding
            finding["severity"] = "hypothesis"
            finding["status"] = "cannot_verify"
            finding["finding"] = f"[HYPOTHESIS — evidence from template, NOT verified in user document] {finding.get('finding', '')}"
            hypothesis_findings.append(finding)
        else:
            verified_findings.append(finding)
    
    # Replace llm_findings with verified only; append hypotheses at end with zero score impact
    llm_findings = verified_findings + hypothesis_findings

    # --- 10.5 Deduplicate redundant findings AND separate hypotheses ---
    # Extract hypotheses; they don't belong in major_findings (separate concern)
    hypothesis_list = [f for f in llm_findings if f.get("severity") == "hypothesis"]
    llm_findings = [f for f in llm_findings if f.get("severity") != "hypothesis"]
    
    # Deduplicate placeholder findings
    placeholder_types = {"template_artifacts", "placeholder_detected"}
    placeholder_findings = [f for f in llm_findings if f.get("type") in placeholder_types]
    non_placeholder = [f for f in llm_findings if f.get("type") not in placeholder_types]
    
    if len(placeholder_findings) > 1:
        # Merge all placeholder findings into one master finding
        total_instances = sum(
            int(re.search(r'(\d+)\s+instance', f.get("finding", "")).group(1))
            for f in placeholder_findings
            if re.search(r'(\d+)\s+instance', f.get("finding", ""))
        ) or 0
        
        all_sections = set()
        all_locations = []
        for f in placeholder_findings:
            loc = f.get("location", "")
            all_locations.append(loc)
            for sec in re.findall(r'\[([^\]]+)\]', loc):
                # Quick filter: skip garbage section names
                if not sec or len(sec) < 4 or sec.startswith(('THEN', 'ELSE', 'IF ', 'Section', 'Sections')):
                    continue
                all_sections.add(sec)
        
        merged_placeholder = {
            "type": "template_artifacts",
            "severity": "warning",
            "location": f"Document-wide ({len(placeholder_findings)} finding types merged)",
            "status": "present",
            "finding": (
                f"Detected {det_stats.get('placeholder_count', 0) + det_stats.get('xxx_count', 0)} total template artifacts "
                f"(placeholders, TBD, XXX markers) across the document. "
                f"Affected sections: {', '.join(sorted(all_sections)[:8])}. "
                f"These should be replaced with finalized content before submission."
            ),
            "why_it_matters": (
                "Template artifacts indicate unfinished content. "
                "While not blocking the overall assessment, they indicate sections that still need "
                "project-specific content. Every placeholder should be replaced before final submission."
            ),
            "evidence": placeholder_findings[0].get("evidence", []),
            "suggested_fix": (
                f"Replace all {det_stats.get('placeholder_count', 0) + det_stats.get('xxx_count', 0)} template artifacts "
                f"with finalized content. Prioritize sections: {', '.join(sorted(all_sections)[:5])}."
            ),
        }
        non_placeholder.insert(0, merged_placeholder)
        llm_findings = non_placeholder

    # --- 11. Merge scores: rubric baseline + LLM adjustments (only from verified findings) ---
    llm_scores = result.get("scores", {})
    final_scores = {}
    for axis in ["structure", "requirements_quality", "traceability",
                  # [COMMENTED OUT] "validation_readiness",  — Validation readiness axis disabled
                  "template_cleanliness", "mechatronics_fitness"]:
        rubric_val = rubric_scores.get(axis, {}).get("score", 0.5)
        llm_val = llm_scores.get(axis, rubric_val)
        if abs(llm_val - rubric_val) < 0.15:
            final_scores[axis] = round(rubric_val, 2)
        else:
            final_scores[axis] = round(0.4 * rubric_val + 0.6 * llm_val, 2)

    # --- 12. Build weighted overall score ---
    placeholder_count = det_stats.get("placeholder_count", 0)
    xxx_count = det_stats.get("xxx_count", 0)
    no_id_count = det_stats.get("no_id_count", 0)
    total_reqs = det_stats.get("total_requirement_rows", 0)
    
    # NOTE: Template artifacts (placeholders) are reported as warnings but do NOT
    # block the overall score. They are informational — the user should fix them,
    # but they don't make the document fundamentally unusable.
    
    # [COMMENTED OUT — Validation readiness weight disabled; redistributed to traceability]
    WEIGHTS = {
        "structure": 0.15,
        "requirements_quality": 0.25,
        "traceability": 0.40,       # Was 0.20; absorbed validation_readiness weight
        # "validation_readiness": 0.20,  # DISABLED — re-enable when validation plans are needed
        "template_cleanliness": 0.10,   # Informational — not blocking
        "mechatronics_fitness": 0.10,
    }
    
    weighted_sum = 0.0
    weight_total = 0.0
    for axis, weight in WEIGHTS.items():
        if axis in final_scores:
            weighted_sum += final_scores[axis] * weight
            weight_total += weight
    
    overall_score = round(weighted_sum / weight_total, 2) if weight_total > 0 else 0.0
    
    # Cap: if many missing IDs, the document has fundamental traceability issues
    if no_id_count > 50 and overall_score > 0.20:
        overall_score = 0.20

    # Build reference lookup for matched_refs
    ref_lookup = {}
    for i, m in enumerate(top_refs, 1):
        ref_lookup[str(i)] = {
            "source_file": m.get("source_file", ""),
            "chunk_id": str(m.get("chunk_id", "")),
            "similarity": m.get("similarity", 0),
        }

    # Build legacy sections with matched_refs populated
    legacy_sections = []
    for finding in llm_findings:
        status_map = {
            "present": "ok",
            "present_but_weak": "warning",
            "present_but_incomplete": "warning",
            "absent": "error",
            "cannot_verify": "warning",
        }
        evidence_refs = []
        for ev in finding.get("evidence", []):
            src_doc = ev.get("source_reference_document", "")
            chunk_id = ev.get("source_section_or_chunk_id", "")
            if src_doc or chunk_id:
                evidence_refs.append(f"{src_doc}::chunk_{chunk_id}" if src_doc else f"chunk_{chunk_id}")
            support = ev.get("support", "")
            ref_matches = re.findall(r'RÉFÉRENCE\s*(\d+)', support)
            for ref_num in ref_matches:
                if ref_num in ref_lookup:
                    rm = ref_lookup[ref_num]
                    evidence_refs.append(f"{rm['source_file']}::chunk_{rm['chunk_id']}")

        legacy_sections.append(SectionFeedback(
            section_name=finding.get("location", "Unknown"),
            status=status_map.get(finding.get("status", ""), "warning"),
            message=finding.get("finding", ""),
            matched_refs=evidence_refs if evidence_refs else [f"retrieved_{len(top_refs)}_chunks"],
        ))

    # --- 12.5 Build per-requirement findings GROUPED by pattern (actionable for engineers) ---
    req_issues: List[RequirementIssue] = []
    placeholder_re = re.compile(r'<<[^>]*>>')
    xxx_re = re.compile(r'\bXXX\b')
    # Detect requirements that reference external standards/documents
    # e.g., "must respect [STA20]", "as specified in [N41]", "per [A7]"
    external_ref_re = re.compile(r'\[[A-Z]+\d*\]')  # [STA20], [N41], [A1]-[A8], etc.
    indirect_req_re = re.compile(
        r'\b(?:respect|comply|accord(?:ance)?|specified|defined|described|refer(?:enced)?|see|per)\b',
        re.IGNORECASE
    )
    
    # Collect all raw issues first
    raw_issues: List[Dict] = []
    for req in requirements:
        req_id = req.get("req_id", "").strip()
        desc = req.get("description", "").strip()
        input_req = req.get("input_requirement", "").strip()
        validation = req.get("validation", "").strip()
        # Compact section: use req_parser's context if available, else empty
        full_sec = req.get("section_context", "") or ""
        section_ctx = _compact_section(full_sec) if full_sec else ""
        location = f"row {req.get('row_index', '?')} in table {req.get('table_index', '?')}"
        
        if not req_id:
            raw_issues.append({"type": "missing_id", "req_id": "[MISSING]", "desc": desc[:150], "loc": location, "sec": section_ctx})
        if placeholder_re.search(desc) or xxx_re.search(desc):
            raw_issues.append({"type": "placeholder_in_desc", "req_id": req_id or "[MISSING]", "desc": desc[:150], "loc": location, "sec": section_ctx})
        if not input_req or input_req.upper() in ("N/A", "N / A", "NA", ""):
            raw_issues.append({"type": "missing_input_ref", "req_id": req_id or "[MISSING]", "desc": desc[:150], "loc": location, "sec": section_ctx})
        # [COMMENTED OUT — Validation method check disabled]
        # if not validation:
        #     raw_issues.append({"type": "missing_validation", "req_id": req_id or "[MISSING]", "desc": desc[:150], "loc": location, "sec": section_ctx})
        if desc and len(desc) < 40:
            raw_issues.append({"type": "weak_description", "req_id": req_id or "[MISSING]", "desc": desc, "loc": location, "sec": section_ctx})
        # Detect requirements that primarily reference external standards/documents
        # Pattern: "must respect [DOC]", "as specified in [DOC]", "comply with [DOC]"
        # These are valid but less verifiable standalone — flagged as WARNING only
        has_external_ref = bool(external_ref_re.search(desc))
        has_indirect_lang = bool(indirect_req_re.search(desc))
        if req_id and has_external_ref and has_indirect_lang:
            # Only flag if the requirement is PRIMARILY a reference (not a concrete spec)
            desc_no_refs = external_ref_re.sub('', desc).strip()
            # If after removing references, little concrete content remains → indirect
            if len(desc_no_refs) < 60:
                raw_issues.append({"type": "indirect_requirement", "req_id": req_id, "desc": desc[:150], "loc": location, "sec": section_ctx})
    
    # Group by issue_type + description pattern (first 3 words)
    groups: Dict[str, List[Dict]] = defaultdict(list)
    for issue in raw_issues:
        # Extract pattern key: first 3 significant words of description
        words = [w for w in issue["desc"].split() if len(w) > 2][:3]
        pattern_key = f"{issue['type']}|{' '.join(words)}" if words else f"{issue['type']}|[empty]"
        groups[pattern_key].append(issue)
    
    # ── Helper functions for requirement issues ─────────────────
    def _get_finding_text(t: str, desc: str) -> str:
        if t == "missing_id": return "Requirement row has no ID — impossible to trace or reference."
        if t == "placeholder_in_desc": return "Description contains placeholder/template artifact."
        if t == "missing_input_ref": return "No upstream/input requirement referenced — breaks traceability chain."
        # [COMMENTED OUT — Validation method finding text disabled]
        # if t == "missing_validation": return "No validation method or acceptance criteria specified."
        if t == "weak_description": return f"Description is very short ({len(desc)} chars) — may lack specificity."
        if t == "indirect_requirement": return (
            "Requirement references an external standard/document instead of providing "
            "directly verifiable criteria. The requirement is valid but less standalone-readable "
            "and less directly testable without consulting the referenced document."
        )
        return "Issue detected."
    
    def _get_fix_text(t: str) -> str:
        if t == "missing_id": return "Assign a unique requirement ID (e.g., REQ-ASU-XXX-NNNN)."
        if t == "placeholder_in_desc": return "Replace placeholder with actual requirement value or specification."
        if t == "missing_input_ref": return "Add reference to the upstream requirement document and ID."
        # [COMMENTED OUT — Validation method fix text disabled]
        # if t == "missing_validation": return "Define test method, acceptance criteria, and conditions."
        if t == "weak_description": return "Expand with preconditions, triggers, observable outcomes, and measurable criteria."
        if t == "indirect_requirement": return (
            "Either (a) add directly verifiable, measurable criteria to the requirement description "
            "so it is self-contained, or (b) accept it as a conformity requirement with the understanding "
            "that the external reference must be available during validation."
        )
        return "Fix the identified issue."
    
    # Emit grouped issues (max 3 per group as examples), skip empty descriptions
    for pattern_key, group_issues in groups.items():
        issue_type = pattern_key.split("|")[0]
        count = len(group_issues)
        
        # Filter out issues with empty descriptions for pattern display
        non_empty = [gi for gi in group_issues if gi["desc"].strip()]
        example_desc = non_empty[0]["desc"][:80] if non_empty else "[empty description]"
        last_desc = non_empty[-1]["desc"][:80] if len(non_empty) > 1 else ""
        
        if count <= 2:
            for gi in group_issues:
                if not gi["desc"].strip():
                    continue  # Skip empty-description issues
                req_issues.append(RequirementIssue(
                    req_id=gi["req_id"],
                    req_description=gi["desc"][:150],
                    issue_type=gi["type"],
                    severity="error" if gi["type"] in ("missing_id", "placeholder_in_desc") else "warning",
                    finding=_get_finding_text(gi["type"], gi["desc"]),
                    suggested_fix=_get_fix_text(gi["type"]),
                    location=gi["loc"],
                    section=gi.get("sec", ""),
                ))
        else:
            # Collect unique sections for grouped issues
            group_sections = list(set(gi.get("sec", "") for gi in group_issues if gi.get("sec", "")))[:3]
            locations = list(set(gi["loc"] for gi in group_issues[:5]))
            req_issues.append(RequirementIssue(
                req_id=f"[{count} requirements]",
                req_description=example_desc[:120],
                issue_type=issue_type,
                severity="error" if issue_type in ("missing_id", "placeholder_in_desc") else "warning",
                finding=f"PATTERN: {count} requirements share this issue. "
                        f"Example: '{example_desc}'"
                        + (f" ... '{last_desc}'" if last_desc else ""),
                suggested_fix=_get_fix_text(issue_type),
                location=f"{len(locations)} locations, e.g. {locations[0]}" if locations else "various",
                section=" | ".join(group_sections) if group_sections else "",
            ))
    
    # Limit to avoid oversized responses
    req_issues = req_issues[:100]

    # --- 12.6 Duplicate requirement detection ---
    # Check if similar requirement descriptions appear in multiple tables
    from difflib import SequenceMatcher
    desc_to_locations: Dict[str, List[str]] = defaultdict(list)
    for req in requirements:
        desc = (req.get("description", "") or "").strip()[:80].lower()
        if len(desc) < 15:
            continue
        loc = f"row {req.get('row_index', '?')} in table {req.get('table_index', '?')}"
        # Normalize for comparison
        normalized = re.sub(r'\s+', ' ', desc)
        desc_to_locations[normalized].append(loc)
    
    # Find duplicates (same or very similar description in different tables)
    duplicate_issues = []
    seen_descs = set()
    for desc, locs in desc_to_locations.items():
        if len(locs) >= 2:
            tables = set(l.split("table ")[-1] for l in locs)
            if len(tables) >= 2 and desc not in seen_descs:
                seen_descs.add(desc)
                duplicate_issues.append({
                    "desc": desc[:100],
                    "locations": locs[:5],
                    "table_count": len(tables),
                })
    
    # Add duplicate findings to major_findings if significant
    if len(duplicate_issues) >= 3:
        dup_finding = {
            "type": "duplicate_requirements",
            "severity": "warning",
            "location": f"{len(duplicate_issues)} requirement patterns",
            "status": "present_but_weak",
            "finding": (
                f"Detected {len(duplicate_issues)} potential duplicate requirements "
                f"appearing in multiple tables. This may indicate copy-paste errors "
                f"or intentional reuse without cross-reference. "
                f"Examples: '{duplicate_issues[0]['desc']}', "
                f"'{duplicate_issues[-1]['desc'] if len(duplicate_issues) > 1 else ''}'"
            ),
            "why_it_matters": (
                "Duplicate requirements across tables risk inconsistency, "
                "double implementation, or conflicting specifications."
            ),
            "evidence": [{
                "source_reference_document": "CTS Writing Guide (duplicate detection)",
                "source_section_or_chunk_id": "deterministic",
                "user_document_excerpt_or_location": f"Found in {duplicate_issues[0]['table_count']} tables: {duplicate_issues[0]['locations']}",
                "support": "Each requirement should appear once with a unique ID. Duplicates must be consolidated or cross-referenced.",
            }],
            "suggested_fix": (
                "Review duplicate requirements. Either consolidate into a single "
                "requirement with a unique ID, or add explicit cross-references "
                "between the duplicated instances."
            ),
        }
        llm_findings.append(dup_finding)

    # --- 13. Honest verdict override based on deterministic evidence ---
    llm_verdict = result.get("global_verdict", "CANNOT_VERIFY")
    placeholder_count = det_stats.get("placeholder_count", 0)
    xxx_count = det_stats.get("xxx_count", 0)
    no_id_count = det_stats.get("no_id_count", 0)
    total_reqs = det_stats.get("total_requirement_rows", 0)
    
    # Determine evidence-based verdict
    # Template artifacts (placeholders) are reported but NOT treated as blocking/critical
    # — they are informational warnings, not verdict-determining issues
    if no_id_count > 50:
        final_verdict = "NON_COMPLIANT"
    elif no_id_count > 30:
        final_verdict = "NOT_RELIABLE"
    elif no_id_count > 15:
        final_verdict = "ACCEPTABLE_WITH_FIXES"
    else:
        final_verdict = llm_verdict if llm_verdict in (
            "GOOD", "ACCEPTABLE_WITH_FIXES", "NOT_RELIABLE", 
            "NON_COMPLIANT", "CANNOT_VERIFY"
        ) else "CANNOT_VERIFY"
    
    # Build honest overall assessment (hypotheses mentioned separately)
    if final_verdict in ("NON_COMPLIANT", "NOT_RELIABLE"):
        honest_assessment = (
            f"DOCUMENT NEEDS SIGNIFICANT IMPROVEMENT. "
            f"Found {placeholder_count} template placeholders (informational — replace before final submission), "
            f"{no_id_count} requirements without IDs, "
            f"{xxx_count} XXX markers. "
            f"Total requirement rows analyzed: {total_reqs}. "
            f"Key issues: missing requirement IDs and incomplete traceability."
            # [COMMENTED OUT — validation methods mention removed]
            # f"Key issues: missing requirement IDs and incomplete validation methods."
            + (f"  ||  NOTE: {len(hypothesis_list)} additional findings are marked as HYPOTHESIS "
               f"(evidence from template, not verified in user document) — these do NOT impact the score."
               if hypothesis_list else "")
        )
    else:
        honest_assessment = result.get("overall_assessment", "")
        if hypothesis_list:
            honest_assessment += (
                f"  ||  NOTE: {len(hypothesis_list)} findings marked as HYPOTHESIS "
                f"(not verified in user document, no score impact)."
            )

    # --- 14. Build section-by-section coverage summary ---
    # Validate section names — filter out parsing artifacts
    GARBAGE_SECTION_PATTERNS = [
        r'^Sections?:?\s*\[', r'^Sections?:?\s',  # "Sections: ...", "Section: ..."
        r'^THEN\s', r'^ELSE\s', r'^IF\s',         # Conditional statements
        r'^OF\s+THE\s', r'^AND\s', r'^OR\s',       # Fragments
        r'^VERSION\s*$', r'^\d+\s', r'^0\s*=?\s',  # Metadata
        r'^\d+$', r'^x+$', r'^\[MISSING\]',         # Numbers/placeholders
        r'^\d+\s+requirements?$',                    # "[14 requirements]"
    ]
    garbage_re = re.compile('|'.join(GARBAGE_SECTION_PATTERNS), re.IGNORECASE)
    
    def _is_valid_section(name: str) -> bool:
        name = name.strip()
        if len(name) < 4 or len(name) > 100:
            return False
        if garbage_re.match(name):
            return False
        if not re.search(r'[A-Za-z]', name):
            return False
        if re.match(r'^[\d\sxX×<>\[\]\(\)\{\}@#$%^&*=+./\\|-]+$', name):
            return False
        return True
    
    # Gather all unique sections mentioned in findings and requirement issues
    all_sections_in_findings: Dict[str, Dict] = {}
    
    # From major findings
    for f in llm_findings:
        loc = f.get("location", "")
        for sec_name in re.findall(r'\[([^\]]+)\]', loc):
            if not _is_valid_section(sec_name):
                continue
            if sec_name not in all_sections_in_findings:
                all_sections_in_findings[sec_name] = {
                    "section": sec_name,
                    "status": "checked_with_issues",
                    "issue_count": 0,
                    "placeholder_count": 0,
                    "requirement_count": 0,
                }
            all_sections_in_findings[sec_name]["issue_count"] += 1
    
    # From requirement issues (group by section)
    for ri in req_issues:
        sec = ri.section
        if sec:
            top_match = re.search(r'\[([^\]]+)\]', sec)
            if top_match:
                top_section = top_match.group(1)
                if not _is_valid_section(top_section):
                    continue
                if top_section not in all_sections_in_findings:
                    all_sections_in_findings[top_section] = {
                        "section": top_section,
                        "status": "checked_with_issues",
                        "issue_count": 0,
                        "placeholder_count": 0,
                        "requirement_count": 0,
                    }
                all_sections_in_findings[top_section]["requirement_count"] += 1
    
    # From placeholder instances
    placeholder_instances = det_stats.get("placeholder_instances", [])
    for pi in placeholder_instances[:50]:
        line_num = pi.get("line", 0)
        sec = get_section_for_line(line_num)
        top_match = re.search(r'\[([^\]]+)\]', sec)
        if top_match:
            top_section = top_match.group(1)
            if not _is_valid_section(top_section):
                continue
            if top_section not in all_sections_in_findings:
                all_sections_in_findings[top_section] = {
                    "section": top_section,
                    "status": "checked_with_issues",
                    "issue_count": 0,
                    "placeholder_count": 0,
                    "requirement_count": 0,
                }
            all_sections_in_findings[top_section]["placeholder_count"] += 1
    
    # Add sections that were checked but had no issues (from the structure check)
    checked_sections = det_stats.get("sections_found", [])
    for sec_name in checked_sections:
        if not _is_valid_section(sec_name):
            continue  # Skip garbage section names from regex false positives
        if sec_name not in all_sections_in_findings:
            all_sections_in_findings[sec_name] = {
                "section": sec_name,
                "status": "checked_ok",
                "issue_count": 0,
                "placeholder_count": 0,
                "requirement_count": 0,
            }
    
    # Merge text-only section analysis results
    for sec_name, analysis in text_section_analysis.items():
        if not _is_valid_section(sec_name):
            continue  # Skip garbage section names
        if sec_name in all_sections_in_findings:
            all_sections_in_findings[sec_name]["text_analysis"] = analysis.get("issues", [])
            if analysis.get("status") == "empty":
                all_sections_in_findings[sec_name]["status"] = "empty_section"
            elif analysis.get("status") == "na_only":
                all_sections_in_findings[sec_name]["status"] = "na_only"
            elif analysis.get("status") == "has_placeholders":
                all_sections_in_findings[sec_name]["status"] = "has_placeholders"
            elif analysis.get("status") == "minimal_content":
                all_sections_in_findings[sec_name]["status"] = "minimal_content"
        else:
            all_sections_in_findings[sec_name] = {
                "section": sec_name,
                "status": analysis.get("status", "checked"),
                "issue_count": len(analysis.get("issues", [])),
                "placeholder_count": 0,
                "requirement_count": 0,
                "text_analysis": analysis.get("issues", []),
            }
    
    # Add reference rules to section summary entries (sanitize angle brackets)
    for sec_name, info in all_sections_in_findings.items():
        rules = get_reference_rules(sec_name, "template")
        if rules:
            # Replace <...> with [...] to prevent HTML stripping in JSON output
            sanitized = [r.replace('<', '[').replace('>', ']') for r in rules[:3]]
            info["applicable_template_rules"] = sanitized
    
    section_summary = sorted(all_sections_in_findings.values(), 
                             key=lambda x: x.get("issue_count", 0) + x.get("placeholder_count", 0), 
                             reverse=True)

    # --- 15. Build HUMAN-READABLE engineering report ---
    human_report = _build_human_report(
        final_verdict, honest_assessment, final_scores, overall_score,
        llm_findings, req_issues, section_summary,
        placeholder_count, no_id_count, xxx_count, total_reqs,
        hypothesis_list
    )
    
    # Cache for GET /report and /report/detailed endpoints
    global _last_human_report, _last_detailed_report, _last_beginner_report
    _last_human_report = human_report

    # Append image analysis section to the human report
    if image_summary.get("total_images_found", 0) > 0:
        img_section = _build_image_report_section(image_results, image_summary)
        _last_human_report = human_report + "\n" + img_section
    
    # Build detailed comparison report (wrapped to prevent 500 on error)
    import traceback
    try:
        detailed_report = build_detailed_report(
            document_name=result.get("document_name", filename or "unknown"),
            verdict=final_verdict,
            overall_score=overall_score,
            scores=final_scores,
            major_findings=llm_findings,
            req_issues=req_issues,
            section_summary=section_summary,
            text_section_analysis=text_section_analysis,
            det_stats=det_stats,
            placeholder_count=placeholder_count,
            no_id_count=no_id_count,
            xxx_count=xxx_count,
            total_reqs=total_reqs,
            hypotheses=hypothesis_list,
            recommendations=result.get("recommendations", []),
            ambiguous_phrases=result.get("ambiguous_phrases", []),
            weak_sections=result.get("weak_sections", []),
        )
    except Exception as e:
        detailed_report = f"Error building detailed report: {str(e)}\n\nTRACEBACK:\n{traceback.format_exc()}"
    _last_detailed_report = detailed_report

    # Build beginner-friendly report (wrapped to prevent 500 on error)
    try:
        beginner_report = build_beginner_report(
            document_name=result.get("document_name", filename or "unknown"),
            verdict=final_verdict,
            overall_score=overall_score,
            scores=final_scores,
            major_findings=llm_findings,
            req_issues=req_issues,
            section_summary=section_summary,
            text_section_analysis=text_section_analysis,
            det_stats=det_stats,
            placeholder_count=placeholder_count,
            no_id_count=no_id_count,
            xxx_count=xxx_count,
            total_reqs=total_reqs,
            hypotheses=hypothesis_list,
            recommendations=result.get("recommendations", []),
            ambiguous_phrases=result.get("ambiguous_phrases", []),
            weak_sections=result.get("weak_sections", []),
        )
    except Exception as e:
        beginner_report = f"Error building beginner report: {str(e)}\n\nTRACEBACK:\n{traceback.format_exc()}"
    _last_beginner_report = beginner_report

    return LeonValidationResponse(
        document_name=result.get("document_name", filename or "unknown"),
        global_verdict=final_verdict,
        overall_assessment=honest_assessment,
        scores=ScoreBreakdown(**final_scores),
        major_findings=[MajorFinding(**f) for f in llm_findings],
        missing_sections=result.get("missing_sections", []) or det_stats.get("missing_cts_sections", []),
        weak_sections=result.get("weak_sections", []),
        ambiguous_phrases=result.get("ambiguous_phrases", []),
        placeholder_or_template_artifacts=result.get("placeholder_or_template_artifacts", []),
        recommendations=result.get("recommendations", []),
        overall_score=overall_score,
        summary=honest_assessment,
        sections=legacy_sections,
        requirement_issues=req_issues,
        section_summary=section_summary,
        human_report=human_report,
        image_analysis=image_summary,
        resolved_standards=resolved_standards,
        standard_codes_detected=standard_codes_in_doc,
    )


@app.get("/report", response_class=Response)
async def get_human_report():
    """
    Retourne le dernier rapport de validation au format TEXTE BRUT (lisible).
    Ouvre directement dans le navigateur avec les sauts de ligne.
    Utilisez cet endpoint après avoir fait un POST /validate.
    """
    if not _last_human_report:
        return Response(
            content="Aucun rapport disponible. Faites d'abord un POST /validate.",
            media_type="text/plain; charset=utf-8"
        )
    return Response(
        content=_last_human_report,
        media_type="text/plain; charset=utf-8"
    )


@app.get("/report/detailed", response_class=Response)
async def get_detailed_report():
    """
    Retourne le rapport DÉTAILLÉ de comparaison des 3 fichiers (Template, Guide, Spec).
    Inclut l'analyse section par section, règle par règle, et toutes les recommandations.
    """
    if not _last_detailed_report:
        return Response(
            content="Aucun rapport détaillé disponible. Faites d'abord un POST /validate.",
            media_type="text/plain; charset=utf-8"
        )
    return Response(
        content=_last_detailed_report,
        media_type="text/plain; charset=utf-8"
    )


@app.get("/report/beginner", response_class=Response)
async def get_beginner_report():
    """
    Retourne un rapport PÉDAGOGIQUE pour débutants — explique chaque section
    en langage simple, compare les 3 fichiers, et donne un plan d'action priorisé.
    IDÉAL pour les ingénieurs AI qui travaillent avec l'équipe mécatronique.
    """
    if not _last_beginner_report:
        return Response(
            content="Aucun rapport débutant disponible. Faites d'abord un POST /validate.",
            media_type="text/plain; charset=utf-8"
        )
    return Response(
        content=_last_beginner_report,
        media_type="text/plain; charset=utf-8"
    )


# ═══════════════════════════════════════════════════════════════════
# beStandard Integration Endpoints
# ═══════════════════════════════════════════════════════════════════

@app.get("/standards/search")
async def search_standards(
    q: str = "",
    code: str = "",
    document_type: str = "",
    status: str = "",
):
    """
    Search the Stellantis beStandard platform for norms/standards.

    Use cases:
    - Look up a standard by code: GET /standards/search?code=STA20
    - Search by title: GET /standards/search?q=acoustic
    - Find all active standards: GET /standards/search?status=PUBLISHED

    Requires BESTANDARD_CLIENT_ID and BESTANDARD_CLIENT_SECRET in .env.
    """
    client = get_bestandard_client()
    if not client.is_configured:
        raise HTTPException(
            status_code=412,
            detail="beStandard not configured. Set BESTANDARD_CLIENT_ID and BESTANDARD_CLIENT_SECRET in .env",
        )

    codes = [code] if code else None
    doc_types = [document_type] if document_type else None
    statuses = [status] if status else None

    try:
        results = client.search(
            query=q,
            codes=codes,
            document_types=doc_types,
            statuses=statuses,
        )
    except RuntimeError as e:
        raise HTTPException(status_code=502, detail=str(e))

    return {
        "query": {"q": q, "code": code, "document_type": document_type, "status": status},
        "total_results": len(results),
        "results": [
            {
                "id": r.id,
                "code": r.code,
                "title": r.title,
                "revision": r.revision,
                "status": r.status,
                "publication_date": r.publication_date,
                "last_modification_date": r.last_modification_date,
            }
            for r in results
        ],
    }


@app.get("/standards/{code}")
async def get_standard_detail(code: str):
    """
    Get full details for a standard by its code (e.g., STA20, N41).

    Returns complete metadata including:
    - Title (EN, IT, FR), revision, status, dates
    - Cross-references (docsInUse, docsRef)
    - Available files (translations, formats)
    - Classification (FCA class, doc type, supplier visibility)

    Requires BESTANDARD_CLIENT_ID and BESTANDARD_CLIENT_SECRET in .env.
    """
    client = get_bestandard_client()
    if not client.is_configured:
        raise HTTPException(
            status_code=412,
            detail="beStandard not configured. Set BESTANDARD_CLIENT_ID and BESTANDARD_CLIENT_SECRET in .env",
        )

    try:
        resolved = client.resolve_standard(code)
    except RuntimeError as e:
        raise HTTPException(status_code=502, detail=str(e))

    if not resolved.found or resolved.norm is None:
        raise HTTPException(
            status_code=404,
            detail=f"Standard '{code}' not found in beStandard: {resolved.error}",
        )

    norm = resolved.norm
    return {
        "code": code,
        "verification_status": resolved.verification_status,
        "detail": {
            "id": norm.id,
            "code": norm.code,
            "title": norm.title,
            "title_it": norm.title_it,
            "title_fr": norm.title_fr,
            "revision": norm.revision,
            "status": norm.status,
            "publication_date": norm.publication_date,
            "review_date": norm.review_date,
            "last_modification_date": norm.last_modification_date,
            "author_dept": norm.author_dept,
            "class_fca": norm.class_fca,
            "doc_type_fca": norm.doc_type_fca,
            "send_to_supplier": norm.send_to_supplier,
            "is_global": norm.is_global,
            "is_cancelled": norm.is_cancelled,
            "is_active": norm.is_active,
            "url_bst": norm.url_bst,
            "cross_references": {
                "docs_in_use": [
                    {"code": d.code, "title": d.title, "status": d.status}
                    for d in norm.docs_in_use
                ],
                "docs_ref": [
                    {"code": d.code, "title": d.title, "status": d.status}
                    for d in norm.docs_ref
                ],
            },
            "files": [
                {
                    "id": f.id,
                    "file_name": f.file_name,
                    "ext": f.ext,
                    "is_published": f.is_published,
                    "is_original": f.is_original,
                    "lang": t.lang,
                    "lang_title": t.title,
                }
                for t in norm.translations
                for f in t.files
            ],
        },
    }


@app.post("/standards/ingest")
async def ingest_standard(code: str = ""):
    """
    Download a standard from beStandard and add it to the dynamic RAG index.

    After ingestion, the validation LLM can verify requirements against
    the actual standard text (not just metadata).

    Usage: POST /standards/ingest?code=STA20

    This enables DEEP VERIFICATION: when a requirement says
    "must comply with STA20 §4.2", LEON can look up what §4.2 says.

    Requires BESTANDARD_CLIENT_ID, BESTANDARD_CLIENT_SECRET, and
    BESTANDARD_DEEP_VERIFY=true in .env.
    """
    if not code:
        raise HTTPException(status_code=400, detail="Parameter 'code' is required")

    try:
        from app.bestandard_ingest import get_standards_index
        std_idx = get_standards_index()
        chunks_added = std_idx.ingest_code(code)
    except RuntimeError as e:
        raise HTTPException(status_code=502, detail=str(e))

    return {
        "code": code,
        "chunks_added": chunks_added,
        "already_indexed": chunks_added == 0,
        "total_indexed_codes": std_idx.get_indexed_codes(),
        "total_index_chunks": len(std_idx._index) if hasattr(std_idx, '_index') else 0,
    }


@app.get("/standards/index/status")
async def get_standards_index_status():
    """
    Get the status of the dynamic standards RAG index.

    Returns which standards are currently indexed and available
    for deep verification during validation.
    """
    try:
        from app.bestandard_ingest import get_standards_index
        std_idx = get_standards_index()
        return {
            "indexed_codes": std_idx.get_indexed_codes(),
            "total_chunks": len(std_idx._index) if hasattr(std_idx, '_index') else 0,
            "index_path": str(std_idx._index if hasattr(std_idx, '_index') else "N/A"),
        }
    except Exception as e:
        return {
            "indexed_codes": [],
            "total_chunks": 0,
            "error": str(e),
        }


def _build_image_report_section(image_results: list, image_summary: dict) -> str:
    """Build a human-readable section about images/diagrams found in the document."""
    lines = []
    lines.append("")
    lines.append("=" * 72)
    lines.append("  IMAGE & DIAGRAM ANALYSIS")
    lines.append("=" * 72)
    lines.append(f"  Total images/diagrams found: {image_summary.get('total_images_found', 0)}")
    lines.append(f"  Sections containing images: {', '.join(image_summary.get('sections_with_images', []))}")
    lines.append(f"  Diagram types detected: {', '.join(image_summary.get('diagram_types', []) or ['unknown'])}")
    lines.append("")
    lines.append("  HOW TO VERIFY: Open your .docx file and look for images/figures in these sections.")
    lines.append("  Each image should have a caption (Figure <n> - <Title>) and be referenced in the text.")
    lines.append("  The CTS template requires that all diagrams follow standard graphic conventions.")
    lines.append("")
    
    for img in image_results[:8]:
        section = img.get("section", "?")
        atype = img.get("analysis_type", "unknown").replace("_", " ")
        desc = img.get("description", "")
        nearby = img.get("detected_from_text", "")
        
        lines.append(f"  ▸ Section: {section}")
        lines.append(f"    Type: {atype}")
        if desc and len(desc) > 10:
            lines.append(f"    Analysis: {desc[:200]}")
        if nearby and nearby != "[No nearby text]":
            lines.append(f"    Nearby text: \"{nearby[:150]}\"")
        lines.append("")
    
    lines.append("  NOTE: For full diagram analysis, the system can be configured to use")
    lines.append("  GPT-4o Vision to analyze diagram content (states, transitions, components).")
    lines.append("  Contact the LEON development team to enable vision-based analysis.")
    
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# beStandard prompt formatters (used in LLM context)
# ---------------------------------------------------------------------------

def _format_resolved_standards_for_prompt(resolved: Dict[str, Dict]) -> str:
    """Format resolved beStandard references for the LLM prompt."""
    if not resolved:
        return "[No external standards were resolved]"
    
    lines = []
    lines.append(f"LEON has resolved {len(resolved)} external standard references via beStandard:")
    lines.append("")
    
    # Sort: resolved/found first, then not found
    sorted_codes = sorted(resolved.keys(), key=lambda c: (not resolved[c].get("found", False), c))
    
    for code in sorted_codes:
        info = resolved[code]
        if info.get("found"):
            status_icon = {
                "active": "✅ ACTIVE",
                "cancelled": "❌ CANCELLED",
                "draft": "⚠️ DRAFT",
            }.get(info.get("verification_status", ""), "❓")
            
            lines.append(
                f"  [{code}] {status_icon} — {info.get('title', 'Unknown')} "
                f"(rev {info.get('revision', '?')}, status: {info.get('status', '?')})"
            )
            if info.get("is_cancelled"):
                lines.append(f"    ⚠️ WARNING: This standard has been CANCELLED. "
                           f"Requirements referencing it should be reviewed.")
        else:
            lines.append(f"  [{code}] ❌ NOT FOUND in beStandard — {info.get('error', 'Unknown error')}")
            lines.append(f"    Requirements referencing [{code}] cannot be verified against this standard.")
    
    lines.append("")
    lines.append("USE THIS INFORMATION: When a requirement references [CODE], check the resolved status above.")
    lines.append("- If ACTIVE: the standard exists, is published, and can be used as a valid reference.")
    lines.append("- If CANCELLED: flag the requirement — it references an obsolete standard.")
    lines.append("- If DRAFT: the standard exists but is not yet finalized — the requirement may change.")
    lines.append("- If NOT FOUND: the code may be a typo, an internal document, or not in beStandard.")
    
    return "\n".join(lines)


def _format_standards_index_for_prompt(doc_text: str, codes: List[str]) -> str:
    """Search the standards RAG index and format results for the LLM prompt."""
    if not codes:
        return "[No standard codes to search]"
    
    try:
        from app.bestandard_ingest import get_standards_index
        std_idx = get_standards_index()
    except Exception:
        return "[Standards index not available]"
    
    indexed = std_idx.get_indexed_codes()
    if not indexed:
        return "[No standards have been ingested into the dynamic index yet. Use POST /standards/ingest to index standards for deep verification.]"
    
    lines = []
    lines.append(f"Standards currently indexed for deep verification: {', '.join(indexed)}")
    lines.append("")
    
    # Search the index for relevant chunks
    # Build a query from requirement descriptions that reference standards
    query_terms = []
    for code in codes:
        # Find requirements that reference this code
        for match in re.finditer(rf'\[{re.escape(code)}\][^.]*\.', doc_text):
            query_terms.append(match.group(0)[:200])
    
    if query_terms:
        combined_query = " ".join(query_terms[:3])[:1000]
        try:
            results = std_idx.search(combined_query, k=5, threshold=0.55, codes=codes)
        except Exception:
            results = []
        
        if results:
            lines.append("Relevant standard excerpts (for verification against requirements):")
            for i, r in enumerate(results, 1):
                lines.append(f"\n  [{r['source_code']} | {r.get('section', '')} | similarity={r['similarity']:.3f}]")
                lines.append(f"  {r['text'][:500]}")
        else:
            lines.append("[No relevant standard excerpts found for the referenced codes. The standards are indexed but no specific sections match the requirements in this document.]")
    else:
        lines.append("[No specific requirement-standard pairs identified for deep verification.]")
    
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Human-readable report builder
# ---------------------------------------------------------------------------

def _build_human_report(
    verdict: str, assessment: str, scores: dict, overall: float,
    findings: list, req_issues: list, section_summary: list,
    placeholder_count: int, no_id_count: int, xxx_count: int, total_reqs: int,
    hypotheses: list
) -> str:
    """Build a human-readable engineering validation report — location-focused and verifiable."""
    
    L = []  # report lines
    
    def sep(char="─", width=72):
        L.append(char * width)
    
    def h1(text):
        L.append(f"\n{'='*72}")
        L.append(f"  {text}")
        L.append(f"{'='*72}")
    
    def h2(text):
        L.append(f"\n  ▸ {text}")
        L.append(f"  {'─'*68}")
    
    def bar(val, w=16):
        f = int(val * w)
        return "█" * f + "░" * (w - f)
    
    # ═══════════════════════════════════════════════════════════
    # HEADER
    # ═══════════════════════════════════════════════════════════
    L.append("╔══════════════════════════════════════════════════════════════════════╗")
    L.append("║  LEON SPEC VALIDATOR — Mechatronics Engineering Validation Report   ║")
    L.append("║  Stellantis Component Technical Specification Audit                 ║")
    L.append("╚══════════════════════════════════════════════════════════════════════╝")
    
    # ═══════════════════════════════════════════════════════════
    # 1. VERDICT
    # ═══════════════════════════════════════════════════════════
    v = {
        "NON_COMPLIANT": "NON-COMPLIANT — Document would be REJECTED at Gate Review",
        "NOT_RELIABLE": "NOT RELIABLE — Major corrections required before use",
        "ACCEPTABLE_WITH_FIXES": "ACCEPTABLE WITH FIXES — Document usable after corrections",
        "GOOD": "GOOD — Document meets Stellantis standards",
    }.get(verdict, verdict)
    
    L.append(f"\n  FINAL VERDICT: {v}")
    L.append(f"  Overall Score: {overall:.2f} / 1.00")
    L.append(f"  {placeholder_count} placeholder artifacts | {no_id_count} requirements without IDs | {xxx_count} XXX markers")
    L.append(f"  {total_reqs} total requirement rows analyzed across the document")
    
    # ═══════════════════════════════════════════════════════════
    # 2. HOW TO READ & VERIFY THIS REPORT
    # ═══════════════════════════════════════════════════════════
    L.append(f"\n  ┌─ HOW TO USE THIS REPORT ─────────────────────────────────────┐")
    L.append(f"  │ For each issue below, you will find:                          │")
    L.append(f"  │   WHERE:    Exact section, table, row, or line number         │")
    L.append(f"  │   PROBLEM:  What is wrong and why it matters                  │")
    L.append(f"  │   FIX:      What to do to resolve the issue                   │")
    L.append(f"  │   VERIFY:   Text to search in your document to confirm        │")
    L.append(f"  │ Open your .docx file and search (Ctrl+F) for the VERIFY text  │")
    L.append(f"  │ to confirm each finding yourself.                             │")
    L.append(f"  └──────────────────────────────────────────────────────────────┘")
    
    # ═══════════════════════════════════════════════════════════
    # 3. SCORE DASHBOARD
    # ═══════════════════════════════════════════════════════════
    L.append(f"\n  QUALITY SCORES (0 = worst, 1 = best)")
    L.append(f"  {'Axis':<30} {'Score':>6}  Bar")
    L.append(f"  {'─'*30} {'─'*6}  {'─'*17}")
    for axis, label in [
        ("structure", "Structure (CTS plan)"),
        ("requirements_quality", "Requirements Quality"),
        ("traceability", "Traceability (upstream links)"),
        ("validation_readiness", "Validation Readiness"),
        ("template_cleanliness", "Template Cleanliness"),
        ("mechatronics_fitness", "Mechatronics Fitness"),
    ]:
        val = scores.get(axis, 0)
        L.append(f"  {label:<30} {val:>5.2f}  {bar(val)}")
    L.append(f"  {'─'*30} {'─'*6}  {'─'*17}")
    L.append(f"  {'OVERALL':<30} {overall:>5.2f}")
    
    # ═══════════════════════════════════════════════════════════
    # 4. CRITICAL ISSUES — WITH EXACT LOCATIONS AND TEXT TO MODIFY
    # ═══════════════════════════════════════════════════════════
    h1("ISSUES FOUND — With Exact Text To Modify")
    
    for i, f in enumerate(findings[:15], 1):
        ftype = f.get("type", "?").replace("_", " ").title()
        loc = f.get("location", "Unknown")
        prob = f.get("finding", "")
        fix = f.get("suggested_fix", "")
        sev = f.get("severity", "info")
        sev_label = "CRITICAL" if sev == "error" else "WARNING" if sev == "warning" else "INFO"
        
        # Extract exact text to modify from evidence (skip metadata-only excerpts)
        doc_excerpt = ""
        evidence_list = f.get("evidence", [])
        if evidence_list:
            ev = evidence_list[0]
            excerpt = ev.get("user_document_excerpt_or_location", "")
            # Skip metadata patterns like "Found in 2 tables: [...]"
            if excerpt.startswith("Found in") or excerpt.startswith("First affected"):
                # For duplicate/metadata findings, use the finding text itself as the reference
                doc_excerpt = ""
            elif "[VERIFIED" in excerpt:
                parts = excerpt.split("] ", 1)
                doc_excerpt = parts[1].strip() if len(parts) > 1 else excerpt.strip()
            elif excerpt and len(excerpt) > 10 and "[WARNING" not in excerpt and "[NEEDS" not in excerpt:
                doc_excerpt = excerpt.strip()
        
        L.append(f"\n  ┌─ ISSUE #{i}: {sev_label} ────────────────────────────────────────┐")
        L.append(f"  │ TYPE:    {ftype}")
        L.append(f"  │ WHERE:   {loc}")
        L.append(f"  │ PROBLEM: {prob[:180]}")
        if doc_excerpt:
            L.append(f"  │ TEXT TO MODIFY (copy this to find in your document):")
            # Show the exact text, truncated for readability
            if len(doc_excerpt) > 150:
                L.append(f"  │   \"{doc_excerpt[:150]}...\"")
            else:
                L.append(f"  │   \"{doc_excerpt}\"")
        if fix:
            L.append(f"  │ FIX:     {fix[:180]}")
        L.append(f"  └{'─'*64}┘")
    
    # ═══════════════════════════════════════════════════════════
    # 5. REQUIREMENT-LEVEL ISSUES — With Exact Text to Modify
    # ═══════════════════════════════════════════════════════════
    h1("REQUIREMENT-LEVEL ISSUES — With Exact Text To Modify")
    L.append(f"  (Each issue shows the exact requirement text that needs fixing)")
    
    # Group requirement issues by section for logical reading
    by_section: dict = {}
    for ri in req_issues:
        sec = ri.section if hasattr(ri, 'section') else ri.get("section", "Unknown")
        if sec not in by_section:
            by_section[sec] = []
        by_section[sec].append(ri)
    
    shown = 0
    for sec_name, issues in sorted(by_section.items()):
        if shown >= 40:
            break
        # Clean up section display: take only the first section if multiple joined by |
        clean_section = sec_name.split(" | ")[0] if " | " in sec_name else sec_name
        L.append(f"\n  ▸ SECTION: {clean_section} ({len(issues)} issue(s))")
        for ri in issues[:5]:
            shown += 1
            rid = ri.req_id if hasattr(ri, 'req_id') else ri.get("req_id", "?")
            # Sanitize req_id: show only first line (the actual ID), strip attribute annotations
            if isinstance(rid, str) and "\n" in rid:
                rid = rid.split("\n")[0].strip()
            itype = (ri.issue_type if hasattr(ri, 'issue_type') else ri.get("issue_type", "?")).replace("_", " ").title()
            loc = ri.location if hasattr(ri, 'location') else ri.get("location", "")
            finding = ri.finding if hasattr(ri, 'finding') else ri.get("finding", "")
            fix = ri.suggested_fix if hasattr(ri, 'suggested_fix') else ri.get("suggested_fix", "")
            desc = ri.req_description if hasattr(ri, 'req_description') else ri.get("req_description", "")
            
            sev_icon = "🔴" if "missing_id" in str(itype).lower() else "🟡"
            L.append(f"     {sev_icon} [{rid}] {itype}")
            L.append(f"        WHERE:        {loc}")
            if desc and desc.strip() and desc.strip() != "[empty description]":
                L.append(f"        TEXT TO FIX:  \"{desc.strip()[:120]}\"")
            elif "[empty description]" in str(desc):
                L.append(f"        TEXT TO FIX:  [Empty cell — needs content]")
            L.append(f"        ISSUE:        {finding[:120]}")
            L.append(f"        FIX:          {fix[:120]}")
    
    if shown >= 40:
        L.append(f"\n  ... and {len(req_issues) - 40} more issues (see JSON for full list)")
    
    # ═══════════════════════════════════════════════════════════
    # 6. SECTION-BY-SECTION STATUS
    # ═══════════════════════════════════════════════════════════
    h1("SECTION-BY-SECTION DOCUMENT STATUS")
    L.append(f"  (Walk through your document section by section to verify)")
    
    issue_secs = [s for s in section_summary if s.get("status") != "checked_ok"]
    ok_secs = [s for s in section_summary if s.get("status") == "checked_ok"]
    
    if issue_secs:
        L.append(f"\n  ▸ SECTIONS REQUIRING ATTENTION:")
        for s in issue_secs[:20]:
            name = s.get("section", "?")
            status = s.get("status", "?").replace("_", " ").title()
            pl = s.get("placeholder_count", 0)
            iss = s.get("issue_count", 0)
            req = s.get("requirement_count", 0)
            
            if "placeholder" in status.lower():
                icon = "🔴"
            elif "minimal" in status.lower() or "na_only" in status.lower():
                icon = "🟡"
            else:
                icon = "🟡"
            
            L.append(f"     {icon} {name}")
            L.append(f"        Status: {status} | Placeholders: {pl} | Issues: {iss} | Requirements: {req}")
            # Show first applicable rule
            rules = s.get("applicable_template_rules", [])
            if rules:
                L.append(f"        Expected: {rules[0][:100]}")
    
    if ok_secs:
        names = [s.get("section", "?") for s in ok_secs[:15]]
        L.append(f"\n  ▸ SECTIONS OK ({len(ok_secs)} total):")
        L.append(f"     🟢 {', '.join(names)}")
    
    # ═══════════════════════════════════════════════════════════
    # 7. PRIORITY ACTION PLAN
    # ═══════════════════════════════════════════════════════════
    h1("PRIORITY ACTION PLAN — What to Fix First")
    
    L.append(f"\n  🔴 CRITICAL (Must fix before Gate Review):")
    L.append(f"     1. Assign unique requirement IDs to {no_id_count} requirements.")
    L.append(f"        Format: REQ-ASU-XXX-NNNN (see CTS template for convention).")
    
    L.append(f"\n  🟡 IMPORTANT (Fix before development use):")
    
    if placeholder_count + xxx_count > 0:
        L.append(f"     2. Replace {placeholder_count + xxx_count} template placeholders with finalized content.")
        L.append(f"        Start with sections: ENVIRONMENT CONDITIONS, DESIGN AND MANUFACTURING, RAMS REQUIREMENTS.")
    
    # [COMMENTED OUT — Validation method counter and action plan item disabled]
    # val_count = sum(1 for ri in req_issues if "validation" in str(ri.issue_type if hasattr(ri, 'issue_type') else ri.get("issue_type", "")))
    trace_count = sum(1 for ri in req_issues if "input_ref" in str(ri.issue_type if hasattr(ri, 'issue_type') else ri.get("issue_type", "")))
    
    # [COMMENTED OUT — Validation method action plan item]
    # if val_count > 0:
    #     L.append(f"     3. Define validation methods for all requirements ({val_count} patterns affected).")
    #     L.append(f"        Each requirement needs: test method + acceptance criteria + test conditions.")
    if trace_count > 0:
        L.append(f"     4. Add upstream requirement references for full traceability ({trace_count} patterns affected).")
        L.append(f"        Fill the 'Input Requirement' column in every requirement table.")
    
    if hypotheses:
        L.append(f"\n  ℹ️  FOR INFORMATION (Hypotheses — not verified, no score impact):")
        for h in hypotheses[:3]:
            L.append(f"     • {h.get('finding', '')[:150]}")
    
    L.append(f"\n  {'='*72}")
    L.append(f"  End of Validation Report")
    L.append(f"  LEON Spec Validator — Stellantis Mechatronics Engineering")
    L.append(f"  {'='*72}")
    
    return "\n".join(L)


# ═════════════════════════════════════════════════════════════════════════════
# LEON Q&A — Posez une question sur un document de spécification
# ═════════════════════════════════════════════════════════════════════════════


@app.post("/ask")
async def ask_question_about_document(
    file: UploadFile = File(..., description="Fichier .docx de la spécification"),
    question: str = Form(..., description="Question sur le document"),
):
    """
    LEON Q&A — Posez une question sur un document de spécification.

    Le système:
    1. Extrait le texte du document .docx
    2. Découpe en chunks et calcule les embeddings
    3. Trouve les passages les plus pertinents pour la question
    4. Génère une réponse UNIQUEMENT à partir de ces passages
       (anti-hallucination: le LLM n'a accès qu'aux extraits du document)

    Retourne:
    - answer: réponse textuelle avec citations de sections
    - citations: liste des passages utilisés avec leur section
    - confidence: high | medium | low | not_found
    - sections_used: sections du document consultées
    """
    from app.ask_leon import ask_question

    # 1. Extraire le texte du DOCX
    temp_path = Path("data/uploads") / (file.filename or "upload_ask.docx")
    temp_path.parent.mkdir(parents=True, exist_ok=True)
    content = await file.read()
    temp_path.write_bytes(content)

    try:
        doc = Document(str(temp_path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        # Extract table cell text
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    ct = cell.text.strip()
                    if ct:
                        table_texts.append(ct)

        document_text = "\n".join(paragraphs + table_texts)

        if not document_text.strip():
            raise HTTPException(status_code=400, detail="Le document est vide")

        # 2. Lancer la Q&A
        result = ask_question(
            document_text=document_text,
            question=question,
            document_name=file.filename or "document",
            top_k=TOP_K_CHUNKS,
            similarity_threshold=SIMILARITY_THRESHOLD,
        )

        if result.error:
            return {
                "answer": result.answer,
                "citations": result.citations,
                "confidence": result.confidence,
                "sections_used": result.sections_used,
                "chunks_retrieved": result.chunks_retrieved,
                "error": result.error,
            }

        return {
            "answer": result.answer,
            "citations": result.citations[:10],  # Top 10 citations
            "confidence": result.confidence,
            "sections_used": result.sections_used,
            "chunks_retrieved": result.chunks_retrieved,
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erreur lors de l'analyse: {e}")
    finally:
        if temp_path.exists():
            temp_path.unlink()


# ---------------------------------------------------------------------------
# Point d'entrée (lancé via uvicorn)
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    import uvicorn
    uvicorn.run("app.main:app", host="0.0.0.0", port=8000, reload=True)
