"""
beStandard RAG Ingestion Pipeline for LEON
═══════════════════════════════════════════════

Downloads standard documents from beStandard, extracts their text,
chunks them, generates embeddings, and builds a dynamic RAG index
that allows the LLM to verify requirements against actual standard text.

This is the DEEP VERIFICATION layer — when a requirement says
"must comply with STA20 §4.2", LEON can now look up what §4.2 actually says.

ARCHITECTURE:
  beStandard API → download PDF/DOCX → extract text → chunk → embed → dynamic index

The dynamic index is separate from the static template/guide index:
  - data/refs/           → static template, writing guide (ingested once)
  - data/standards_index/ → dynamic standards content (ingested on-demand)

USAGE:
  from app.bestandard_ingest import ingest_standard_by_code, StandardsIndex

  idx = StandardsIndex()
  idx.ingest_code("STA20")  # Downloads STA20, chunks it, embeds it
  results = idx.search("acoustic output level requirements", k=5)
"""

import json
import logging
import tempfile
import os
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, field

from app.config import (
    DATA_DIR,
    EMBEDDING_MODEL_MAX_INPUT,
)
from app.bestandard_client import (
    BeStandardClient,
    get_bestandard_client,
    NormDetail,
    ResolvedStandard,
)
from app.embeddings import get_embedding, find_similar_chunks, cosine_similarity

logger = logging.getLogger(__name__)

# Where the dynamic standards index lives
STANDARDS_INDEX_DIR: Path = DATA_DIR / "standards_index"
STANDARDS_INDEX_PATH: Path = STANDARDS_INDEX_DIR / "standards_index.json"


# ═══════════════════════════════════════════════════════════════════
# TEXT EXTRACTION (from downloaded standard files)
# ═══════════════════════════════════════════════════════════════════

def _extract_text_from_pdf_bytes(pdf_bytes: bytes) -> str:
    """
    Extract text from PDF binary content.

    Uses PyPDF2 if available, falls back to a simple byte reader.
    """
    try:
        from PyPDF2 import PdfReader
        import io
        reader = PdfReader(io.BytesIO(pdf_bytes))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n\n".join(pages)
    except ImportError:
        logger.warning("PyPDF2 not installed — cannot extract PDF text. Install with: pip install PyPDF2")
        return ""
    except Exception as e:
        logger.warning("PDF extraction failed: %s", e)
        return ""


def _extract_text_from_docx_bytes(docx_bytes: bytes) -> str:
    """
    Extract text from DOCX binary content.
    """
    try:
        from docx import Document
        import io
        doc = Document(io.BytesIO(docx_bytes))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # Also extract table text
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    ct = cell.text.strip()
                    if ct:
                        table_texts.append(ct)
        return "\n".join(paragraphs + table_texts)
    except Exception as e:
        logger.warning("DOCX extraction failed: %s", e)
        return ""


def extract_text_from_file_bytes(file_bytes: bytes, file_ext: str) -> str:
    """
    Route to the right extractor based on file extension.

    Args:
        file_bytes: Raw file binary content
        file_ext: File extension ("pdf", "docx", "txt", etc.)

    Returns:
        Extracted text string (empty if extraction fails)
    """
    ext = file_ext.lower().lstrip('.')
    if ext == "pdf":
        return _extract_text_from_pdf_bytes(file_bytes)
    elif ext in ("docx", "doc"):
        return _extract_text_from_docx_bytes(file_bytes)
    elif ext == "txt":
        try:
            return file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            try:
                return file_bytes.decode('latin-1')
            except Exception:
                return ""
    else:
        logger.warning("Unsupported file extension: .%s", ext)
        return ""


# ═══════════════════════════════════════════════════════════════════
# CHUNKING (simple but effective for standards)
# ═══════════════════════════════════════════════════════════════════

def _chunk_standard_text(text: str, source_code: str, chunk_size: int = 800) -> List[Dict[str, Any]]:
    """
    Chunk a standard document into overlapping text segments.

    Standards are typically structured: section numbers, requirements, tables.
    We use a paragraph-based approach with overlap to preserve context.

    Args:
        text: Raw extracted text
        source_code: The standard code (e.g., "STA20") for metadata
        chunk_size: Target characters per chunk (before overlap)

    Returns:
        List of chunk dicts with text + metadata
    """
    import re

    chunks = []
    paragraphs = text.split('\n')
    paragraphs = [p.strip() for p in paragraphs if p.strip()]

    # Try to detect section boundaries (e.g., "1.", "1.1", "§4.2", "4.2.1")
    section_re = re.compile(r'^(?:§\s*)?(\d+(?:\.\d+)*)\s+')

    current_chunk: List[str] = []
    current_size = 0
    current_section = ""
    chunk_id = 0

    def _flush():
        nonlocal chunk_id
        if current_chunk:
            chunk_text = " ".join(current_chunk)
            chunks.append({
                "text": chunk_text,
                "source_file": f"bestandard://{source_code}",
                "chunk_id": chunk_id,
                "section_context": current_section or source_code,
                "source_code": source_code,
            })
            chunk_id += 1
            current_chunk.clear()

    for para in paragraphs:
        # Detect section headers
        m = section_re.match(para)
        if m:
            current_section = f"{source_code} §{m.group(1)}"
            # Flush previous chunk on section boundary
            if current_size > chunk_size * 0.5:
                _flush()
                current_size = 0

        para_len = len(para)

        if current_size + para_len > chunk_size and current_chunk:
            _flush()
            current_size = 0

        current_chunk.append(para)
        current_size += para_len

    _flush()  # Don't forget the last chunk
    return chunks


# ═══════════════════════════════════════════════════════════════════
# STANDARDS INDEX
# ═══════════════════════════════════════════════════════════════════

@dataclass
class StandardIndexEntry:
    """An indexed chunk from a standard document."""
    source_code: str          # e.g., "STA20"
    source_title: str         # e.g., "Acoustic Vehicle Alerting System"
    text: str                 # The chunk text
    embedding: List[float]    # Vector embedding
    section: str = ""         # Section within the standard
    chunk_id: int = 0         # Chunk number


class StandardsIndex:
    """
    Dynamic RAG index for standards downloaded from beStandard.

    Unlike the static template index (built once from data/refs/),
    this index grows on-demand as new standards are resolved.

    Persisted to data/standards_index/standards_index.json.
    """

    def __init__(self):
        self._client: Optional[BeStandardClient] = None
        self._index: List[StandardIndexEntry] = []
        self._loaded = False

    @property
    def client(self) -> BeStandardClient:
        if self._client is None:
            self._client = get_bestandard_client()
        return self._client

    def _ensure_loaded(self):
        """Lazy-load the index from disk."""
        if self._loaded:
            return
        self._loaded = True

        if not STANDARDS_INDEX_PATH.exists():
            return

        try:
            with open(STANDARDS_INDEX_PATH, 'r', encoding='utf-8') as f:
                raw = json.load(f)
            for entry in raw:
                self._index.append(StandardIndexEntry(**entry))
            logger.info(
                "StandardsIndex: loaded %d chunks from %s",
                len(self._index), STANDARDS_INDEX_PATH
            )
        except Exception as e:
            logger.warning("StandardsIndex: failed to load: %s", e)

    def _save(self):
        """Persist the index to disk."""
        STANDARDS_INDEX_DIR.mkdir(parents=True, exist_ok=True)
        try:
            raw = [
                {
                    "source_code": e.source_code,
                    "source_title": e.source_title,
                    "text": e.text,
                    "embedding": e.embedding,
                    "section": e.section,
                    "chunk_id": e.chunk_id,
                }
                for e in self._index
            ]
            with open(STANDARDS_INDEX_PATH, 'w', encoding='utf-8') as f:
                json.dump(raw, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.warning("StandardsIndex: failed to save: %s", e)

    def is_indexed(self, code: str) -> bool:
        """Check if a standard code is already in the index."""
        self._ensure_loaded()
        return any(e.source_code.upper() == code.upper() for e in self._index)

    def ingest_code(self, code: str) -> int:
        """
        Download a standard from beStandard and add it to the RAG index.

        Args:
            code: Standard code, e.g. "STA20"

        Returns:
            Number of chunks added to the index

        Raises:
            RuntimeError: If the standard cannot be resolved or downloaded
        """
        self._ensure_loaded()

        # Don't re-ingest
        if self.is_indexed(code):
            logger.info("StandardsIndex: %s already indexed, skipping.", code)
            return 0

        if not self.client.is_configured:
            raise RuntimeError("beStandard not configured — cannot ingest standards")

        # Resolve the standard
        resolved = self.client.resolve_standard(code)
        if not resolved.found or resolved.norm is None:
            raise RuntimeError(f"Standard '{code}' not found in beStandard: {resolved.error}")

        norm = resolved.norm
        logger.info(
            "StandardsIndex: ingesting %s — %s (status: %s)",
            code, norm.title, norm.status
        )

        # Download the first published file
        file_bytes = self.client.download_first_published_file(norm)
        if file_bytes is None:
            raise RuntimeError(f"No published files for standard '{code}'")

        # Determine file extension from the published file
        published_files = norm.published_files
        file_ext = published_files[0].ext if published_files else "pdf"

        # Extract text
        text = extract_text_from_file_bytes(file_bytes, file_ext)
        if not text or len(text) < 100:
            raise RuntimeError(
                f"Failed to extract text from standard '{code}' "
                f"(ext={file_ext}, bytes={len(file_bytes)})"
            )

        logger.info(
            "StandardsIndex: extracted %d chars from %s (.%s)",
            len(text), code, file_ext
        )

        # Chunk
        chunks = _chunk_standard_text(text, code, chunk_size=800)
        logger.info("StandardsIndex: %d chunks to embed", len(chunks))

        # Embed each chunk
        new_entries = 0
        for chunk in chunks:
            try:
                embedding = get_embedding(chunk["text"][:EMBEDDING_MODEL_MAX_INPUT])
            except Exception as e:
                logger.warning("StandardsIndex: embedding failed for chunk %d: %s", chunk["chunk_id"], e)
                continue

            self._index.append(StandardIndexEntry(
                source_code=code,
                source_title=norm.title,
                text=chunk["text"],
                embedding=embedding,
                section=chunk.get("section_context", ""),
                chunk_id=chunk["chunk_id"],
            ))
            new_entries += 1

        # Persist
        self._save()
        logger.info(
            "StandardsIndex: ingested %s — %d chunks added (total: %d chunks)",
            code, new_entries, len(self._index)
        )
        return new_entries

    def search(
        self,
        query: str,
        k: int = 5,
        threshold: float = 0.65,
        codes: Optional[List[str]] = None,
    ) -> List[Dict[str, Any]]:
        """
        Search the standards index for relevant chunks.

        Args:
            query: Search query text
            k: Number of top results to return
            threshold: Minimum cosine similarity
            codes: Optional filter — only search specific standard codes

        Returns:
            List of matching chunks with similarity scores
        """
        self._ensure_loaded()

        if not self._index:
            return []

        try:
            query_embedding = get_embedding(query[:EMBEDDING_MODEL_MAX_INPUT])
        except Exception as e:
            logger.warning("StandardsIndex: query embedding failed: %s", e)
            return []

        results = []
        for entry in self._index:
            # Optional code filter
            if codes and entry.source_code.upper() not in [c.upper() for c in codes]:
                continue

            sim = cosine_similarity(query_embedding, entry.embedding)
            if sim >= threshold:
                results.append({
                    "source_code": entry.source_code,
                    "source_title": entry.source_title,
                    "section": entry.section,
                    "text": entry.text,
                    "similarity": round(sim, 4),
                    "chunk_id": entry.chunk_id,
                })

        results.sort(key=lambda x: x["similarity"], reverse=True)
        return results[:k]

    def get_indexed_codes(self) -> List[str]:
        """Get list of standard codes currently in the index."""
        self._ensure_loaded()
        return sorted(set(e.source_code for e in self._index))

    def clear(self):
        """Clear the in-memory and on-disk index."""
        self._index.clear()
        if STANDARDS_INDEX_PATH.exists():
            STANDARDS_INDEX_PATH.unlink()
        logger.info("StandardsIndex: cleared")


# ═══════════════════════════════════════════════════════════════════
# SINGLETON
# ═══════════════════════════════════════════════════════════════════

_index: Optional[StandardsIndex] = None


def get_standards_index() -> StandardsIndex:
    """Get or create the singleton standards index."""
    global _index
    if _index is None:
        _index = StandardsIndex()
    return _index
