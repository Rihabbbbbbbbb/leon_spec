"""
Retrieval logic for the Q&A assistant.

Indexes accessible specification files (plain-text extracts under data/) and
retrieves the most relevant passages for a user question.

Design:
- Accessible files are discovered from a configurable directory (default: data/).
- Each file is split into overlapping chunks.
- Retrieval uses cosine similarity over Azure OpenAI embeddings when available.
- A lightweight keyword fallback is used when embeddings are not configured,
  so the assistant can run before the real retrieval is connected.
- BeStandard / standards repositories are NEVER used here.
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

from app.config import DATA_DIR, UPLOADS_DIR, SIMILARITY_THRESHOLD, TOP_K_CHUNKS


# ── Data structures ────────────────────────────────────────────────
@dataclass
class Chunk:
    """A retrievable passage from an accessible spec file."""
    file_name: str
    text: str
    section: str = ""
    chunk_id: int = 0


@dataclass
class RetrievalResult:
    """Result of a retrieval call."""
    chunks: List[Chunk] = field(default_factory=list)
    scores: List[float] = field(default_factory=list)
    used_fallback: bool = False
    error: Optional[str] = None


# ── Chunking ───────────────────────────────────────────────────────
_SENTENCE_END = re.compile(r"(?<=[.!?])\s+")
_CHUNK_CHAR_TARGET = 900   # ~220 tokens
_CHUNK_CHAR_OVERLAP = 180

# Headings commonly found in CTS/ASU specs (uppercase or numbered sections).
_HEADING_RE = re.compile(
    r"^(?:#+\s*)?("
    r"[A-Z][A-Z\s/()\-&,:;.\u2013\u2014]{3,}"      # ALLCAPS SECTION
    r"|\d+(?:\.\d+)*\.?\s+[A-Z][A-Za-z\s/()\-&,:;.]{3,}"  # 1.2 Title
    r")\s*$",
    re.MULTILINE,
)


def _detect_section(line: str) -> str:
    """Return a clean section name if the line looks like a heading, else ''."""
    m = _HEADING_RE.match(line.strip())
    if not m:
        return ""
    name = m.group(1).strip().rstrip(":")
    return name


def _split_into_chunks(text: str, file_name: str) -> List[Chunk]:
    """Split raw text into section-aware overlapping chunks for retrieval."""
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if not text:
        return []

    # Track current section as we walk lines, so each chunk carries context.
    lines = text.split("\n")
    current_section = ""
    paragraphs: List[tuple] = []  # (section, paragraph_text)
    para: List[str] = []

    for line in lines:
        stripped = line.strip()
        sec = _detect_section(stripped)
        if sec and len(stripped) < 80:
            # Flush current paragraph under the previous section
            if para:
                paragraphs.append((current_section, " ".join(para)))
                para = []
            current_section = sec
            continue
        if not stripped:
            if para:
                paragraphs.append((current_section, " ".join(para)))
                para = []
            continue
        para.append(stripped)
    if para:
        paragraphs.append((current_section, " ".join(para)))

    # Group paragraphs into ~target-size chunks, preserving section context.
    chunks: List[Chunk] = []
    buf: List[str] = []
    buf_section = ""
    buf_len = 0
    chunk_id = 0

    for section, ptext in paragraphs:
        if buf_len + len(ptext) > _CHUNK_CHAR_TARGET and buf:
            chunks.append(Chunk(
                file_name=file_name,
                text=" ".join(buf),
                section=buf_section,
                chunk_id=chunk_id,
            ))
            overlap: List[str] = []
            overlap_len = 0
            for s in reversed(buf):
                if overlap_len + len(s) > _CHUNK_CHAR_OVERLAP:
                    break
                overlap.insert(0, s)
                overlap_len += len(s)
            buf = overlap
            buf_len = sum(len(s) for s in buf)
            chunk_id += 1
            buf_section = section
        if not buf:
            buf_section = section
        buf.append(ptext)
        buf_len += len(ptext)

    if buf:
        chunks.append(Chunk(
            file_name=file_name,
            text=" ".join(buf),
            section=buf_section,
            chunk_id=chunk_id,
        ))
    return chunks


# ── Accessible file discovery ──────────────────────────────────────
_DEFAULT_SPEC_FILES = [
    "spec_extracted.txt",
    "template_extracted.txt",
]


def discover_accessible_files(root: Optional[Path] = None) -> List[Path]:
    """
    Return the list of accessible specification files.

    Includes:
    - built-in spec extracts under data/
    - user-uploaded files under data/uploads/ AND UPLOADS_DIR (Azure /tmp)
    BeStandard / standards repositories are intentionally excluded.
    """
    root = root or DATA_DIR
    files: List[Path] = []
    for name in _DEFAULT_SPEC_FILES:
        p = root / name
        if p.exists():
            files.append(p)
    # Also include any *.txt and *.docx under data/refs/ if present
    refs = root / "refs"
    if refs.exists():
        files.extend(sorted(refs.glob("*.txt")))
        files.extend(sorted(refs.glob("*.docx")))
    # Include user-uploaded spec files from UPLOADS_DIR
    # (In Azure Functions, UPLOADS_DIR may differ from root/"uploads")
    uploaded = list_uploaded_files()
    # Avoid duplicates if UPLOADS_DIR == root/"uploads"
    existing_names = {p.name for p in files}
    for p in uploaded:
        if p.name not in existing_names:
            files.append(p)
            existing_names.add(p.name)
    return files


def build_index(root: Optional[Path] = None) -> List[Chunk]:
    """Read all accessible spec files and return a list of chunks."""
    chunks: List[Chunk] = []
    for path in discover_accessible_files(root):
        # Use the proper extractor for DOCX/PDF; read TXT directly.
        if path.suffix.lower() == ".txt":
            try:
                text = path.read_text(encoding="utf-8", errors="replace")
            except OSError:
                continue
        else:
            text = extract_text_from_file(path)
        if not text or not text.strip():
            continue
        chunks.extend(_split_into_chunks(text, path.name))
    return chunks


# ── Uploaded file text extraction ──────────────────────────────────
def extract_text_from_file(path: Path) -> str:
    """
    Extract plain text from an uploaded specification file.

    Supports:
    - .txt  : read directly
    - .docx : structured extraction via python-docx (paragraphs + tables)
    - .pdf  : extraction via PyPDF2 if available
    Returns empty string on failure.
    """
    path = Path(path)
    suffix = path.suffix.lower()
    try:
        if suffix == ".txt":
            return path.read_text(encoding="utf-8", errors="replace")
        if suffix == ".docx":
            return _extract_docx(path)
        if suffix == ".pdf":
            return _extract_pdf(path)
    except Exception:
        return ""
    return ""


def _extract_docx(path: Path) -> str:
    """Extract text from a .docx file (paragraphs + table cells)."""
    from docx import Document
    doc = Document(str(path))
    parts: List[str] = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_pdf(path: Path) -> str:
    """Extract text from a .pdf file via PyPDF2 (optional dependency)."""
    try:
        from PyPDF2 import PdfReader
    except Exception:
        return ""
    reader = PdfReader(str(path))
    parts: List[str] = []
    for page in reader.pages:
        t = (page.extract_text() or "").strip()
        if t:
            parts.append(t)
    return "\n".join(parts)


# ── Uploaded files index management ────────────────────────────────
def save_uploaded_file(file_name: str, content: bytes) -> Path:
    """Save uploaded file bytes to the uploads directory and return its path."""
    UPLOADS_DIR.mkdir(parents=True, exist_ok=True)
    # Sanitize the file name
    safe = re.sub(r"[^A-Za-z0-9._\- ]", "_", file_name).strip()
    if not safe:
        safe = "uploaded_spec.txt"
    dest = UPLOADS_DIR / safe
    dest.write_bytes(content)
    return dest


def list_uploaded_files() -> List[Path]:
    """Return uploaded spec files saved in the uploads directory."""
    if not UPLOADS_DIR.exists():
        return []
    exts = {".txt", ".docx", ".pdf"}
    return sorted(p for p in UPLOADS_DIR.iterdir() if p.suffix.lower() in exts)


def build_uploaded_index() -> List[Chunk]:
    """Build a chunk index from uploaded specification files only."""
    chunks: List[Chunk] = []
    for path in list_uploaded_files():
        text = extract_text_from_file(path)
        if text:
            chunks.extend(_split_into_chunks(text, path.name))
    return chunks


def delete_uploaded_file(file_name: str) -> bool:
    """Delete an uploaded file by name. Returns True if deleted."""
    safe = re.sub(r"[^A-Za-z0-9._\- ]", "_", file_name).strip()
    target = UPLOADS_DIR / safe
    if target.exists():
        target.unlink()
        return True
    return False


# ── Keyword fallback retrieval ─────────────────────────────────────
_STOP = {
    "the", "a", "an", "is", "are", "was", "were", "be", "been", "being",
    "what", "who", "where", "when", "why", "how", "which", "whose",
    "of", "to", "in", "on", "for", "and", "or", "with", "by", "as", "at",
    "from", "into", "this", "that", "these", "those", "it", "its",
    "does", "do", "did", "can", "could", "should", "would", "shall",
    "will", "may", "might", "must", "have", "has", "had",
    "please", "tell", "me", "about", "there", "their", "they",
    "le", "la", "les", "de", "des", "du", "et", "que", "quoi",
    "comment", "pourquoi", "où", "quand", "qui", "quel", "quelle",
    "est", "sont", "dans", "sur", "pour", "avec",
}


def _tokenize(text: str) -> set:
    # Match alphanumeric + accented characters (French: éèêëàâäùûüïîôöç)
    # \w with re.UNICODE matches accented chars, but we also strip apostrophes
    # inside words (l'ASU → lasu) before tokenizing for better matching.
    text = text.replace("'", " ").replace("\u2019", " ")  # apostrophes
    return set(re.findall(r"[a-zàâäéèêëïîôöùûüçñ0-9]+", text.lower()))


# ── French → English translation map for cross-language retrieval ──
# Engineering specs are often in English; users may ask in French.
# We expand French query terms to their English equivalents so keyword
# retrieval still finds the right passages. This is NOT translation of
# the answer — only of the query tokens used for matching.
_FR_EN_MAP = {
    # core nouns
    "composant": "component", "composants": "components",
    "système": "system", "systeme": "system", "systèmes": "systems",
    "specification": "specification", "spécification": "specification",
    "document": "document", "fichier": "file",
    "exigence": "requirement", "exigences": "requirements",
    "fonction": "function", "fonctions": "functions",
    "fonctionnel": "functional", "fonctionnelle": "functional",
    "performance": "performance", "contrainte": "constraint",
    "interface": "interface", "interfaces": "interfaces",
    "validation": "validation", "test": "test", "tests": "tests",
    "bruit": "noise", "force": "force", "couple": "torque",
    "température": "temperature", "temperature": "temperature",
    "vitesse": "speed", "poids": "weight", "masse": "mass",
    "tension": "voltage", "courant": "current",
    "matériau": "material", "materiau": "material",
    "sécurité": "safety", "securite": "safety",
    "fiabilité": "reliability", "fiabilite": "reliability",
    "durée": "duration", "duree": "duration", "vie": "life",
    "cycle": "cycle", "cycles": "cycles",
    "porte": "door", "poignée": "handle", "poignee": "handle",
    "verrou": "lock", "verrouillage": "locking",
    "clé": "key", "cle": "key", "cylindre": "cylinder",
    "corrosion": "corrosion", "étanchéité": "sealing", "etancheite": "sealing",
    "objectif": "purpose", "but": "purpose", "rôle": "role", "role": "role",
    "description": "description", "description": "description",
    "emplacement": "location", "localisation": "location",
    "schéma": "diagram", "schema": "diagram",
    "norme": "standard", "standard": "standard",
    "conformité": "compliance", "conformite": "compliance",
    "qualité": "quality", "qualite": "quality",
    "erreur": "error", "défaut": "fault", "defaut": "fault",
    "signal": "signal", "communication": "communication",
    "alimentation": "power", "énergie": "energy", "energie": "energy",
    "montage": "mounting", "installation": "installation",
    "dimension": "dimension", "dimensions": "dimensions",
    "tolérance": "tolerance", "tolerance": "tolerance",
    # verbs / question words
    "explique": "explain", "expliquer": "explain",
    "décris": "describe", "decris": "describe", "décrire": "describe",
    "donne": "give", "lister": "list", "liste": "list",
    "quelle": "what", "quel": "what", "quelles": "what", "quels": "what",
    "combien": "how many", "pourquoi": "why",
    "où": "where", "quand": "when", "comment": "how",
    "cherche": "find", "trouver": "find",
    "vérifie": "verify", "verifie": "verify", "vérifier": "verify",
    # adjectives
    "maximal": "maximum", "maximale": "maximum", "maximum": "maximum",
    "minimal": "minimum", "minimale": "minimum", "minimum": "minimum",
    "électrique": "electrical", "electrique": "electrical",
    "mécanique": "mechanical", "mecanique": "mechanical",
    # location / position (critical for "où est situé" questions)
    "situé": "located", "situe": "located", "située": "located",
    "situés": "located", "situées": "located",
    "situ": "located",  # without accent (fallback)
    "position": "position", "positions": "positions",
    "emplacement": "location",
    "capot": "hood", "plenum": "plenum",
    "arrière": "rear", "arriere": "rear",
    "avant": "front",
    "gauche": "left", "droite": "right",
    "véhicule": "vehicle", "vehicule": "vehicle",
    "où": "where",  # with accent
    # alarm / security terms
    "alarme": "alarm", "alerte": "alert",
    "intrusion": "intrusion", "vol": "theft",
    "son": "sound", "bruit": "noise",
    "détecter": "detect", "detecter": "detect",
    "surveillance": "surveillance",
    # states / modes
    "état": "state", "etat": "state", "états": "states",
    "mode": "mode", "modes": "modes",
    "actif": "active", "active": "active",
    "désactivé": "deactivated", "desactive": "deactivated",
    # more verbs
    "fonctionne": "works", "fonctionnent": "works",
    "marche": "works",
    "indique": "indicates", "indiquer": "indicates",
    "produit": "produces", "produire": "produce",
    "émet": "emits", "emet": "emits", "émettre": "emit",
    "reçoit": "receives", "recoit": "receives", "recevoir": "receive",
    # connectors
    "est": "is", "sont": "are",
    "dans": "in", "sur": "on", "sous": "under",
    "entre": "between", "vers": "towards",
    "ce": "this", "cette": "this", "ces": "these",
    "le": "the", "la": "the", "les": "the",
    "un": "a", "une": "a",
    "du": "of the", "des": "of the",
    "de": "of",
    "et": "and",
    "ou": "or",
    "pas": "not", "non": "not",
    "avec": "with", "sans": "without",
    "pour": "for",
    "par": "by",
    "à": "to", "a": "to",
}


def _expand_french_tokens(tokens: set) -> set:
    """Expand French tokens to their English equivalents for cross-language retrieval."""
    expanded = set(tokens)
    for t in list(tokens):
        en = _FR_EN_MAP.get(t)
        if en:
            # Add the English equivalent (and its sub-tokens if multi-word)
            for sub in en.split():
                expanded.add(sub)
    return expanded


def _content_tokens(text: str) -> set:
    """Tokens with stop words removed, plus French→English expansion."""
    raw = _tokenize(text) - _STOP
    return _expand_french_tokens(raw)


def keyword_retrieve(
    question: str,
    chunks: List[Chunk],
    top_k: int = TOP_K_CHUNKS,
    min_score: float = 0.10,
) -> RetrievalResult:
    """
    Lightweight keyword retrieval with TF-IDF-like weighting.

    Score combines:
    - query coverage: fraction of query content tokens found in the chunk
    - chunk specificity: penalize chunks that match only because they're huge
    - phrase bonus: reward adjacent query-word co-occurrence
    """
    q_tokens = _content_tokens(question)
    if not q_tokens:
        return RetrievalResult(used_fallback=True)

    # Document frequency for IDF weighting
    n_chunks = len(chunks)
    df: dict = {}
    chunk_token_sets: List[set] = []
    for ch in chunks:
        toks = _content_tokens(ch.text)
        chunk_token_sets.append(toks)
        for t in toks:
            df[t] = df.get(t, 0) + 1

    import math
    scored: List[tuple] = []
    for ch, toks in zip(chunks, chunk_token_sets):
        overlap = q_tokens & toks
        # Also check section-name overlap: if the chunk's section name
        # contains query tokens, include the chunk even if the text doesn't
        sec_overlap = set()
        if ch.section:
            sec_tokens = set(_tokenize(ch.section.lower()))
            sec_overlap = q_tokens & sec_tokens
        if not overlap and not sec_overlap:
            continue
        # IDF-weighted coverage of the query
        idf_sum = sum(math.log(1 + n_chunks / max(df.get(t, 1), 1)) for t in overlap)
        coverage = idf_sum / (1.0 + len(q_tokens))
        # Specificity: prefer shorter chunks that still match (denser signal)
        specificity = len(overlap) / (1.0 + math.log(1 + len(toks)))
        # Phrase bonus: adjacent query tokens appearing close in text
        low = ch.text.lower()
        phrase_bonus = 0.0
        q_list = sorted(q_tokens)
        for i in range(len(q_list) - 1):
            if q_list[i] in low and q_list[i + 1] in low:
                p1 = low.find(q_list[i])
                p2 = low.find(q_list[i + 1])
                if 0 <= abs(p1 - p2) <= 40:
                    phrase_bonus += 0.15
        score = coverage + 0.3 * specificity + phrase_bonus
        # Section-name bonus: if the chunk's section name contains query
        # tokens, boost the score significantly — this helps questions like
        # "what are the external interfaces?" match the EXTERNAL INTERFACES
        # section even when the chunk text doesn't contain those exact words.
        if sec_overlap:
                # Strong bonus: section name is a very strong signal of
                # topical relevance. 2.0 per matching token ensures section-
                # name matches can compete with text matches.
                section_bonus = 2.0 * len(sec_overlap) / max(len(q_tokens), 1)
                score += section_bonus
        # Source priority: uploaded files (non-built-in) get a small boost
        # so the user's uploaded spec is preferred over built-in extracts.
        _BUILTIN_FILES = {"spec_extracted.txt", "template_extracted.txt"}
        if ch.file_name not in _BUILTIN_FILES:
            score *= 1.15  # 15% boost for uploaded files
        if score >= min_score:
            scored.append((score, ch))
    scored.sort(key=lambda x: x[0], reverse=True)

    top = scored[:top_k]
    return RetrievalResult(
        chunks=[c for _, c in top],
        scores=[s for s, _ in top],
        used_fallback=True,
    )


# ── Semantic retrieval (Azure OpenAI embeddings) ───────────────────
def semantic_retrieve(
    question: str,
    chunks: List[Chunk],
    top_k: int = TOP_K_CHUNKS,
    threshold: float = SIMILARITY_THRESHOLD,
) -> RetrievalResult:
    """Embedding-based retrieval using Azure OpenAI."""
    try:
        from app.embeddings import get_embedding, cosine_similarity  # local import
    except Exception as exc:  # pragma: no cover
        return RetrievalResult(error=f"embeddings module unavailable: {exc}")

    try:
        q_emb = get_embedding(question)
    except Exception as exc:
        return RetrievalResult(error=f"embedding call failed: {exc}")

    scored: List[tuple] = []
    for ch in chunks:
        try:
            c_emb = get_embedding(ch.text[:8000])
        except Exception:
            continue
        score = cosine_similarity(q_emb, c_emb)
        if score >= threshold:
            scored.append((score, ch))
    scored.sort(key=lambda x: x[0], reverse=True)

    top = scored[:top_k]
    return RetrievalResult(
        chunks=[c for _, c in top],
        scores=[s for s, _ in top],
        used_fallback=False,
    )


# ── Azure AI Search hybrid retrieval ───────────────────────────────
def azure_search_retrieve(
    question: str,
    top_k: int = TOP_K_CHUNKS,
    filter_expr: str = "",
) -> RetrievalResult:
    """
    Production retrieval via Azure AI Search (hybrid vector + full-text).

    Falls back gracefully to keyword_retrieve() if:
    - Azure Search is not configured (no endpoint/key)
    - Azure Search is unreachable (network error)
    - The index is empty (no documents)

    This is the PREFERRED retrieval path for the Azure Function.
    It eliminates cold-start penalties (3-5s with FAISS) and provides
    true semantic understanding via pre-computed embedding vectors.
    """
    import logging

    # ── Check if Azure Search is configured ────────────────────────
    try:
        from app.qa.azure_search import is_configured, hybrid_search
    except ImportError:
        logging.info("azure_search module not available; using keyword fallback")
        return keyword_retrieve(question, build_index(), top_k=top_k)

    if not is_configured():
        logging.info("Azure Search not configured; using keyword fallback")
        return keyword_retrieve(question, build_index(), top_k=top_k)

    # ── Get query embedding ────────────────────────────────────────
    try:
        from app.embeddings import get_embedding
        query_embedding = get_embedding(question)
    except Exception as exc:
        logging.warning(f"Query embedding failed: {exc}; using keyword fallback")
        return keyword_retrieve(question, build_index(), top_k=top_k)

    # ── Execute hybrid search ──────────────────────────────────────
    try:
        filter_expr_clean = filter_expr if filter_expr else None
        docs = hybrid_search(
            question=question,
            query_embedding=query_embedding,
            top_k=top_k,
            filter_expr=filter_expr_clean,
        )
    except Exception as exc:
        logging.warning(f"Azure Search query failed: {exc}; using keyword fallback")
        return keyword_retrieve(question, build_index(), top_k=top_k)

    # ── Convert results to RetrievalResult ─────────────────────────
    if not docs:
        # No results from Azure Search — try keyword fallback as safety net
        logging.info("Azure Search returned 0 results; trying keyword fallback")
        return keyword_retrieve(question, build_index(), top_k=top_k)

    chunks = []
    scores = []
    for doc in docs:
        chunks.append(Chunk(
            file_name=doc.get("file_name", ""),
            text=doc.get("text", ""),
            section=doc.get("section", ""),
            chunk_id=doc.get("chunk_id", 0),
        ))
        # Normalize Azure Search score to 0-1 range for compatibility
        # @search.score is unbounded; cap at reasonable range
        raw_score = doc.get("@search.score", 0.0)
        scores.append(min(raw_score / 10.0, 1.0))

    return RetrievalResult(
        chunks=chunks,
        scores=scores,
        used_fallback=False,
    )


# ── Public entry point ─────────────────────────────────────────────
def retrieve(
    question: str,
    chunks: Optional[List[Chunk]] = None,
    top_k: int = TOP_K_CHUNKS,
    use_semantic: bool = False,
) -> RetrievalResult:
    """
    Retrieve relevant passages for a question.

    Args:
        question: user question
        chunks: pre-built index; built on the fly if None
        top_k: number of passages to return
        use_semantic: if True, try Azure OpenAI embeddings; fall back to keywords
    """
    if chunks is None:
        chunks = build_index()
    if not chunks:
        return RetrievalResult(error="no accessible specification files indexed")

    if use_semantic:
        result = semantic_retrieve(question, chunks, top_k=top_k)
        if result.chunks or result.error:
            return result
        # fall through to keyword fallback if semantic returned nothing
    return keyword_retrieve(question, chunks, top_k=top_k)
